import streamlit as st
import google.generativeai as genai
import tempfile
import sounddevice as sd
import numpy as np
import wave
import threading
import time
import os
from pathlib import Path
import base64
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx2pdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black, HexColor
from io import BytesIO
import mysql.connector
from mysql.connector import Error
import uuid
import json

# Replace with your actual API key
Google_API_Key = "AIzaSyCw8tLyHeobBO65GGnkLUVCGSMLdg-HsBw"  # Replace with your valid API key
genai.configure(api_key=Google_API_Key)
model = genai.GenerativeModel('gemini-1.5-flash')

# Set page configuration for a professional look
st.set_page_config(
    page_title="RESUWHISPER AI",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# MySQL Connection Function
def create_connection():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            database='resuwhisper',
            user='resuwhisper_user',
            password='sachinbhat'
        )
        if connection.is_connected():
            return connection
    except Error as e:
        st.error(f"Error connecting to MySQL: {e}")
        return None

# MySQL User Authentication Functions
def signup(username, password):
    connection = create_connection()
    if connection:
        cursor = connection.cursor()
        try:
            cursor.execute("INSERT INTO users (username, password) VALUES (%s, %s)", (username, password))
            connection.commit()
            return True
        except Error as e:
            st.error(f"Error during signup: {e}")
            return False
        finally:
            cursor.close()
            connection.close()
    return False

def login(username, password):
    connection = create_connection()
    if connection:
        cursor = connection.cursor(dictionary=True)
        try:
            cursor.execute("SELECT * FROM users WHERE username = %s AND password = %s", (username, password))
            user = cursor.fetchone()
            if user:
                return user['id']
            else:
                return None
        except Error as e:
            st.error(f"Error during login: {e}")
            return None
        finally:
            cursor.close()
            connection.close()
    return None

# MySQL Session Management Functions
def start_new_session(user_id):
    session_id = str(uuid.uuid4())
    connection = create_connection()
    if connection:
        cursor = connection.cursor()
        try:
            cursor.execute("INSERT INTO user_sessions (user_id, session_id) VALUES (%s, %s)", (user_id, session_id))
            connection.commit()
            return session_id
        except Error as e:
            st.error(f"Error starting new session: {e}")
            return None
        finally:
            cursor.close()
            connection.close()
    return None

def update_session_data(session_id, field, value):
    connection = create_connection()
    if connection:
        cursor = connection.cursor()
        try:
            if field in ['responses', 'resume_data']:
                cursor.execute(f"UPDATE user_sessions SET {field} = %s WHERE session_id = %s", (json.dumps(value), session_id))
            elif field == 'final_resume':
                cursor.execute("UPDATE user_sessions SET final_resume = %s WHERE session_id = %s", (value, session_id))
            else:
                cursor.execute(f"UPDATE user_sessions SET {field} = %s WHERE session_id = %s", (value, session_id))
            connection.commit()
        except Error as e:
            st.error(f"Error updating session data: {e}")
        finally:
            cursor.close()
            connection.close()

def get_session_data(session_id):
    connection = create_connection()
    if connection:
        cursor = connection.cursor(dictionary=True)
        try:
            cursor.execute("SELECT * FROM user_sessions WHERE session_id = %s", (session_id,))
            session = cursor.fetchone()
            if session:
                return session
            else:
                return None
        except Error as e:
            st.error(f"Error retrieving session data: {e}")
            return None
        finally:
            cursor.close()
            connection.close()
    return None

# Resume templates and other constants
resume_templates = {
    "Fresher": "Ideal for 0-2 years of experience. Maximum 1 page.",
    "Intermediate": "Best for 3-7 years of experience. Maximum 1 page.",
    "Veteran": "For 7+ years of experience. Maximum of 2 pages."
}

languages = [
    "English", "Hindi", "Konkani", "Kannada", "Dogri", "Bodo", "Urdu", "Tamil",
    "Kashmiri", "Assamese", "Bengali", "Marathi", "Sindhi", "Maithili",
    "Punjabi", "Malayalam", "Manipuri", "Telugu", "Sanskrit", "Nepali",
    "Santali", "Gujarati", "Odia"
]

questions = [
    "What is your full name, age, address, phone number, email, and LinkedIn or GitHub profile link (if any)?",
    "What are your career goals, key strengths, and professional personality traits?",
    "What is your work experience? Include job titles, company names, employment dates, responsibilities, and achievements.",
    "Tell us about the projects you have done. Include project names, descriptions, and your contributions (optional).",
    "What is your educational background? Include degrees, institutions, graduation years, and any relevant coursework or honors.",
    "What are your hard skills (e.g., technical skills) and soft skills (e.g., communication, teamwork)?",
    "What certifications do you have? Include the certification name, issuing organization, and date received.",
    "What are your extracurricular activities or recognitions? Include activities, organizations, dates, and achievements."
]

section_headers = [
    "📋 Personal Information",
    "💼 Professional Summary",
    "👔 Work Experience",
    "🚀 Projects",
    "🎓 Education",
    "🛠️ Skills",
    "🏆 Certifications",
    "🌟 Extracurricular Activities"
]

# Remove in-memory users_db initialization since we're using MySQL
# if "users_db" not in st.session_state:
#     st.session_state["users_db"] = {}  # Format: {username: {"password": password, "data": resume_data}}

# Initialize session states
def init_session_state():
    defaults = {
        "page": "login",
        "authenticated": False,
        "username": None,
        "session_id": None,  # Add session_id to track the current session
        "selected_language": None,
        "consent_given": False,
        "resume_template": None,
        "current_question_index": 0,
        "responses": {},
        "recording_state": False,
        "audio_file": None,
        "stop_event": None,
        "current_response": None,
        "resume_data": {
            "personal_info": {
                "full_name": "",
                "degree": "",
                "phone": "",
                "email": "",
                "linkedin": "",
                "github": "",
                "address": ""
            },
            "summary": "",
            "qualifications": [],
            "certifications": [],
            "skills": [],
            "experience": [],
            "projects": [],
            "positions": []
        },
        "transcribed_once": {},  # To track if an upload has been transcribed
        "translated_questions": {}  # To store translated questions
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

def record_audio(filename, stop_event: threading.Event, samplerate=44100):
    audio_data = []
    
    def callback(indata, frames, time, status):
        if not stop_event.is_set():
            audio_data.append(indata.copy())
        else:
            raise sd.StopStream()

    try:
        with sd.InputStream(samplerate=samplerate, channels=1, dtype=np.int16, callback=callback):
            st.write("🎤 Recording... Press 'Stop Recording' to end.")
            while not stop_event.is_set():
                time.sleep(0.1)

        st.write("✅ Recording Completed.")
        audio_data = np.concatenate(audio_data, axis=0)

        with wave.open(filename, 'wb') as wf:
            wf.setnchannels(1)
            wf.setsampwidth(2)
            wf.setframerate(samplerate)
            wf.writeframes(audio_data.tobytes())

        if os.path.exists(filename) and os.path.getsize(filename) > 0:
            st.success(f"✅ Recording saved to {filename}!")
        else:
            st.error("Audio file was not created or is empty")
    except sd.StopStream:
        pass
    except Exception as e:
        st.error(f"Recording error: {str(e)}")

def get_gemini_response(input_msg, audio_path=None, mime_type="audio/wav"):
    try:
        if audio_path:
            with open(audio_path, "rb") as f:
                audio = genai.upload_file(audio_path, mime_type=mime_type)
            response = model.generate_content([audio, input_msg])
        else:
            response = model.generate_content([input_msg])
        return response.text if response else None
    except Exception as e:
        st.error(f"Error with Gemini: {str(e)}")
        return None

def translate_questions(language):
    if language not in st.session_state["translated_questions"]:
        prompt = f"""
        Translate the following 8 English questions into {language}. Provide the output as a numbered list, with each translation corresponding to the original question. Do not add extra commentary or modify the questions beyond translation.

        1. What is your full name, age, address, phone number, email, and LinkedIn or GitHub profile link (if any)?
        2. What are your career goals, key strengths, and professional personality traits?
        3. What is your work experience? Include job titles, company names, employment dates, responsibilities, and achievements.
        4. Tell us about the projects you have done. Include project names, descriptions, and your contributions (optional).
        5. What is your educational background? Include degrees, institutions, graduation years, and any relevant coursework or honors.
        6. What are your hard skills (e.g., technical skills) and soft skills (e.g., communication, teamwork)?
        7. What certifications do you have? Include the certification name, issuing organization, and date received.
        8. What are your extracurricular activities or recognitions? Include activities, organizations, dates, and achievements.
        """
        response = get_gemini_response(prompt)
        if response:
            translated = [line.split(". ", 1)[1] for line in response.strip().split("\n") if line.strip()]
            st.session_state["translated_questions"][language] = translated[:8]  # Ensure exactly 8 questions
        else:
            st.session_state["translated_questions"][language] = questions  # Fallback to English if translation fails
    return st.session_state["translated_questions"][language]

def process_response_with_gemini(question_index, response):
    prompts = [
        f"""
        **Your Goal**: Your goal is to meticulously extract and enhance personal details from the user's response to create a flawless, professional resume entry, ensuring every detail is accurate and perfectly formatted for the live resume editor.
        **Your Role**: You are a professional resume maker with years of experience crafting polished, error-free resumes for clients across industries. Make sure you be as accurate as possible.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to extract only what the user explicitly states—no assumptions, no inventions—and enhance it with professional formatting.
        **Detailed Instructions**:
        - Extract the following fields only if explicitly mentioned in the response:
          - Full Name (e.g., "John Doe")
          - Age (e.g., "22")
          - Address (e.g., "123 Main St, Springfield, IL, USA")
          - Phone Number (e.g., "+1-123-456-7890")
          - Email Address (e.g., "john.doe@example.com")
          - LinkedIn Profile (e.g., "linkedin.com/in/johndoe")
          - GitHub Profile (e.g., "github.com/johndoe")
        - Do not guess or add information the user didn’t provide. If a field is missing, leave it blank (e.g., "Age: ").
        - Enhance formatting for professionalism:
          - Standardize phone numbers (e.g., convert "1234567890" to "+1-123-456-7890").
          - Ensure email addresses are lowercase and valid (e.g., "John@EXAMPLE.com" becomes "john@example.com").
          - Verify URLs are complete (e.g., add "https://" to "linkedin.com/in/johndoe" if missing).
        - Avoid duplicating or misplacing fields (e.g., don’t put email in the address field).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        - Put what sounds like an "email" in the email field. What sounds like an address must be put int the address field. 
        **Example**:
        - User Input: "My name is John Doe, I live at 123 Main St, email is john@example, phone 123-456-7890, linkedin.com/in/johndoe"
        - Output:
          Full Name: John Doe
          Age: 
          Address: 123 Main St
          Phone Number: +1-123-456-7890
          Email Address: john@example.com
          LinkedIn Profile: https://linkedin.com/in/johndoe
          GitHub Profile: 
        **Output Format**: Return the extracted and enhanced data in this exact format, one field per line:
        Full Name: [Name]
        Age: [Age]
        Address: [Address]
        Phone Number: [Phone]
        Email Address: [Email]
        LinkedIn Profile: [LinkedIn]
        GitHub Profile: [GitHub]
        """,
        f"""
        **Your Goal**: Your goal is to craft a concise, impactful, and professional summary for the user's resume based solely on their response, enhancing it to sound polished and tailored for the live resume editor without adding unmentioned details.
        **Your Role**: You are a professional resume maker skilled at transforming raw career aspirations and traits into compelling summaries that impress hiring managers.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to enhance only what the user explicitly mentions into a 50-70 word summary.
        **Detailed Instructions**:
        - Focus on three key areas if mentioned:
          - Career Goals (e.g., "Aspiring software developer")
          - Key Strengths (e.g., "problem-solving, coding")
          - Professional Traits (e.g., "team player, detail-oriented")
        - Do not invent details or add fluff. If the user doesn’t mention a category, omit it.
        - Enhance the language to be professional and concise:
          - Convert casual phrases (e.g., "I like coding") to polished ones (e.g., "Skilled in coding").
          - Combine elements logically (e.g., "I want to be a developer and I’m good at teamwork" becomes "Aspiring developer with strong teamwork skills").
        - Keep the summary between 50-70 words for brevity and impact.
        - Do not include suggestions or additional commentary—output only the enhanced summary.
        **Example**:
        - User Input: "I want to be a software engineer, I’m good at problem-solving and I work well in teams."
        - Output:
          Enhanced Summary: Aspiring software engineer with a strong aptitude for problem-solving and a proven ability to collaborate effectively in team environments. Passionate about leveraging technical skills to develop innovative solutions and contribute to organizational success.
        **Output Format**: Return the enhanced summary in this exact format:
        Enhanced Summary: [Summary]
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s work experience into detailed, structured fields for the live resume editor, ensuring each job is accurately separated and professionally formatted without errors or omissions.
        **Your Role**: You are a professional resume maker with expertise in organizingLOW work history into clear, impactful entries that highlight achievements and responsibilities.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to break down multiple jobs if mentioned and enhance only what’s provided.
        **Detailed Instructions**:
        - For each job mentioned, extract these fields if explicitly stated:
          - Job Title (e.g., "Software Intern")
          - Company Name (e.g., "Tech Corp")
          - Dates (e.g., "June 2022 - August 2022")
          - Responsibilities (e.g., "Developed features using Python")
          - Achievements (e.g., "Improved performance by 20%")
        - Identify multiple jobs by keywords like "and," "then," "also," or distinct company names. Separate each job with "---".
        - Do not assume details not provided. Leave fields blank if missing (e.g., "Achievements: ").
        - Enhance professionalism:
          - Use action verbs (e.g., "Coded" becomes "Developed").
          - Quantify where possible if numbers are explicitly stated (e.g., "Improved performance by 20%").
        - Ensure no overlap or misplacement (e.g., dates don’t go into responsibilities).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I was a Software Intern at Tech Corp from June 2022 to August 2022, coded stuff, and then a Developer at XYZ Inc from Jan 2023 to now, built apps."
        - Output:
          Job Title: Software Intern
          Company Name: Tech Corp
          Dates: June 2022 - August 2022
          Responsibilities: Developed software features
          Achievements: 
          ---
          Job Title: Developer
          Company Name: XYZ Inc
          Dates: Jan 2023 - Present
          Responsibilities: Developed mobile applications
          Achievements: 
        **Output Format**: Return each job in this exact format, separated by "---":
        Job Title: [Title]
        Company Name: [Company]
        Dates: [Dates]
        Responsibilities: [Responsibilities]
        Achievements: [Achievements]
        ---
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s project details into concise, professional entries for the live resume editor, treating this as an optional section and ensuring accuracy without adding unmentioned projects.
        **Your Role**: You are a professional resume maker adept at showcasing projects to highlight technical skills and contributions in a clear, employer-friendly format.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each project separately if multiple are mentioned.
        **Detailed Instructions**:
        - For each project mentioned, extract:
          - Project Name (e.g., "Portfolio Website")
          - Description (e.g., "Built a responsive site")
          - Contribution (e.g., "Coded frontend in React")
        - Identify multiple projects by keywords like "and," "also," or distinct names. List each with a "-".
        - Do not add projects not mentioned. If none are provided, return an empty response.
        - Enhance for clarity and professionalism:
          - Combine details into a single line (e.g., "Made a site" becomes "Portfolio Website: Built a responsive site, coded frontend").
          - Use technical terms where applicable (e.g., "Helped on app" becomes "Mobile App: Assisted in development, contributed to UI").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I did a Portfolio Website, built a site, and a Chat App, coded it with friends."
        - Output:
          - Portfolio Website: Built a responsive website, Coded frontend
          - Chat App: Developed a real-time application, Collaborated on coding
        **Output Format**: Return each project in this exact format, one per line:
        - [Project Name]: [Description], [Contribution]
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s educational background into detailed, accurate entries for the live resume editor, capturing every institution mentioned and formatting them professionally without missing or misplacing any details.
        **Your Role**: You are a professional resume maker with a keen eye for organizing educational history into a clear, chronological format that appeals to employers.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each educational entry separately.
        **Detailed Instructions**:
        - For each educational institution mentioned, extract:
          - Degree (e.g., "B.Tech in Computer Science")
          - Institution (e.g., "XYZ University")
          - Graduation Year (e.g., "2023")
          - Honors (e.g., "Cum Laude", if mentioned)
        - Identify multiple entries by keywords like "and," "also," "then," or distinct institutions. List each with a "-".
        - Do not add unmentioned details or fields (e.g., no "Relevant Coursework" unless stated).
        - Enhance for professionalism:
          - Standardize degree names (e.g., "CS degree" becomes "B.Tech in Computer Science").
          - Ensure years are four digits (e.g., "23" becomes "2023").
        - Order entries chronologically if dates are provided, most recent first.
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "B.Tech from IIT in 2023, Diploma from XYZ in 2020, and high school from ABC in 2018."
        - Output:
          - B.Tech in Computer Science, IIT, 2023
          - Diploma, XYZ, 2020
          - High School Diploma, ABC, 2018
        **Output Format**: Return each education entry in this exact format, one per line:
        - [Degree], [Institution], [Graduation Year], [Honors if mentioned]
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s skills into a clear, categorized list for the live resume editor, ensuring every skill is accurately placed and professionally presented without errors or omissions.
        **Your Role**: You are a professional resume maker skilled at identifying and classifying skills to maximize their impact on a resume.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to separate hard and soft skills explicitly.
        **Detailed Instructions**:
        - Categorize skills into two groups if mentioned:
          - Hard Skills (technical, e.g., "Python, SQL, Machine Learning")
          - Soft Skills (interpersonal, e.g., "Communication, Teamwork")
        - Extract only skills the user states. Do not add unmentioned skills.
        - Enhance for clarity:
          - List skills as a comma-separated string (e.g., "coding, teamwork" becomes "Hard Skills: Python, Soft Skills: Teamwork").
          - Specify languages/tools for vague terms where explicitly stated (e.g., "Coding" becomes "Python" if Python is mentioned).
        - Avoid duplication or misclassification (e.g., "Teamwork" stays in soft skills, not hard).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I know Python, teamwork, and problem-solving."
        - Output:
          - Hard Skills: Python
          - Soft Skills: Teamwork, Problem-Solving
        **Output Format**: Return skills in this exact format, one category per line:
        - Hard Skills: [Hard Skills]
        - Soft Skills: [Soft Skills]
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s certifications into precise, professional entries for the live resume editor, ensuring every certification is captured accurately and formatted consistently without errors.
        **Your Role**: You are a professional resume maker experienced in presenting certifications to boost a candidate’s credibility.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each certification separately.
        **Detailed Instructions**:
        - For each certification mentioned, extract:
          - Certification Name (e.g., "AWS Certified Developer")
          - Issuing Organization (e.g., "Amazon")
          - Date (e.g., "July 2023")
        - Identify multiple certifications by keywords like "and," "also," or distinct names. List each with a "-".
        - Do not add unmentioned certifications. If none are provided, return an empty response.
        - Enhance for professionalism:
          - Standardize names (e.g., "AWS cert" becomes "AWS Certified Developer").
          - Format dates as "Month Year" (e.g., "7/23" becomes "July 2023").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "AWS cert from Amazon in July 2023, and Python cert from Coursera in 2022."
        - Output:
          - AWS Certified Developer, Amazon, July 2023
          - Python Certification, Coursera, 2022
        **Output Format**: Return each certification in this exact format, one per line:
        - [Certification Name], [Issuing Organization], [Date]
        """,
        f"""
        **Your Goal**: Your goal is to extract and enhance the user’s extracurricular activities into detailed, professional entries for the live resume editor, treating this as optional and ensuring accuracy without adding unmentioned details.
        **Your Role**: You are a professional resume maker skilled at highlighting extracurriculars to showcase leadership and soft skills.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each activity separately if multiple are mentioned.
        **Detailed Instructions**:
        - For each activity mentioned, extract:
          - Title (e.g., "President of Coding Club")
          - Organization (e.g., "XYZ University")
          - Dates (e.g., "2022-2023")
          - Accomplishment (e.g., "Organized hackathon")
        - Identify multiple activities by keywords like "and," "also," or distinct titles. List each with a "-".
        - Do not add unmentioned activities. If none are provided, return an empty response.
        - Enhance for professionalism:
          - Use active language (e.g., "Helped club" becomes "Organized club events").
          - Standardize dates (e.g., "22-23" becomes "2022-2023").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "President of Coding Club at XYZ University 2022-2023, organized hackathon, and volunteer at ABC Org 2021."
        - Output:
          - President of Coding Club, XYZ University, 2022-2023, Organized hackathon
          - Volunteer, ABC Org, 2021
        **Output Format**: Return each activity in this exact format, one per line:
        - [Title], [Organization], [Dates], [Accomplishment]
        """
    ]

    prompt = prompts[question_index]
    result = get_gemini_response(prompt)
    return result

def update_resume_data(question_index, response):
    result = process_response_with_gemini(question_index, response)
    if not result:
        st.error("Failed to process response with Gemini")
        return

    if question_index == 0:
        lines = result.split("\n")
        personal_info = st.session_state["resume_data"]["personal_info"]
        for line in lines:
            if "Full Name:" in line:
                personal_info["full_name"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Address:" in line:
                personal_info["address"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Phone Number:" in line:
                personal_info["phone"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Email Address:" in line:
                personal_info["email"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "LinkedIn Profile:" in line:
                personal_info["linkedin"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "GitHub Profile:" in line:
                personal_info["github"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
        st.session_state["resume_data"]["personal_info"] = personal_info

    elif question_index == 1:
        for line in result.split("\n"):
            if "Enhanced Summary:" in line:
                summary = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
                st.session_state["resume_data"]["summary"] = summary
                break

    elif question_index == 2:
        experience = []
        current_exp = {}
        for line in result.split("\n"):
            if line.strip() == "---" and current_exp:
                experience.append(current_exp)
                current_exp = {}
            elif "Job Title:" in line:
                current_exp["job_title"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Company Name:" in line:
                current_exp["company"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Dates:" in line:
                current_exp["dates"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Responsibilities:" in line:
                current_exp["responsibilities"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
            elif "Achievements:" in line:
                current_exp["achievements"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""
        if current_exp:
            experience.append(current_exp)
        st.session_state["resume_data"]["experience"] = experience

    elif question_index == 3:
        projects = []
        for line in result.split("\n"):
            if line.startswith("-"):
                projects.append(line.strip("- ").strip())
        st.session_state["resume_data"]["projects"] = projects

    elif question_index == 4:
        qualifications = []
        for line in result.split("\n"):
            if line.startswith("-"):
                qualifications.append(line.strip("- ").strip())
        st.session_state["resume_data"]["qualifications"] = qualifications

    elif question_index == 5:
        skills = []
        for line in result.split("\n"):
            if line.startswith("-"):
                skills.append(line.strip("- ").strip())
        st.session_state["resume_data"]["skills"] = skills

    elif question_index == 6:
        certifications = []
        for line in result.split("\n"):
            if line.startswith("-"):
                certifications.append(line.strip("- ").strip())
        st.session_state["resume_data"]["certifications"] = certifications

    elif question_index == 7:
        positions = []
        for line in result.split("\n"):
            if line.startswith("-"):
                positions.append(line.strip("- ").strip())
        st.session_state["resume_data"]["positions"] = positions

    # Update resume_data in the database
    if st.session_state["session_id"]:
        update_session_data(st.session_state["session_id"], "resume_data", st.session_state["resume_data"])

def generate_word_resume(return_pdf=False):
    resume_data = st.session_state["resume_data"]
    doc = DocxDocument()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

    header = doc.add_paragraph()
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.add_run(resume_data["personal_info"]["full_name"] or "Full Name Not Provided")
    run.bold = True
    run.font.size = Pt(16)
    
    if resume_data["personal_info"]["degree"]:
        run = header.add_run(f"\n{resume_data['personal_info']['degree']}")
    else:
        run = header.add_run("\nDegree Not Provided")
    run.bold = True
    run.font.size = Pt(12)
    
    contact_info_parts = []
    if resume_data["personal_info"]["phone"]:
        contact_info_parts.append(resume_data["personal_info"]["phone"])
    else:
        contact_info_parts.append("Phone Not Provided")
    if resume_data["personal_info"]["email"]:
        contact_info_parts.append(resume_data["personal_info"]["email"])
    else:
        contact_info_parts.append("Email Not Provided")
    if resume_data["personal_info"]["linkedin"]:
        contact_info_parts.append(resume_data["personal_info"]["linkedin"])
    if resume_data["personal_info"]["github"]:
        contact_info_parts.append(resume_data["personal_info"]["github"])
    contact_info = " | ".join(contact_info_parts)
    run = header.add_run(f"\n{contact_info}")
    run.font.size = Pt(10)
    
    run = header.add_run(f"\n{resume_data['personal_info']['address'] or 'Address Not Provided'}")
    run.font.size = Pt(10)

    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph(resume_data["summary"] or "Summary Not Provided")
    p.style.font.size = Pt(10)

    doc.add_heading("Experience", level=1)
    if resume_data["experience"]:
        for exp in resume_data["experience"]:
            p = doc.add_paragraph(f"{exp.get('job_title', 'Title Not Provided')} at {exp.get('company', 'Company Not Provided')} ({exp.get('dates', 'Dates Not Provided')})")
            p.style.font.size = Pt(10)
            p.runs[0].bold = True
            if exp.get("responsibilities"):
                p = doc.add_paragraph(exp["responsibilities"], style='List Bullet')
                p.style.font.size = Pt(10)
            if exp.get("achievements"):
                p = doc.add_paragraph(exp["achievements"], style='List Bullet')
                p.style.font.size = Pt(10)
    else:
        p = doc.add_paragraph("Experience Not Provided")
        p.style.font.size = Pt(10)

    if resume_data["projects"]:
        doc.add_heading("Projects", level=1)
        for proj in resume_data["projects"]:
            p = doc.add_paragraph(proj, style='List Bullet')
            p.style.font.size = Pt(10)

    doc.add_heading("Qualifications & Certifications", level=1)
    all_qualifications = resume_data["qualifications"] + resume_data["certifications"]
    if all_qualifications:
        for qual in all_qualifications:
            if qual:  # Only include non-empty entries
                p = doc.add_paragraph(qual, style='List Bullet')
                p.style.font.size = Pt(10)
    else:
        p = doc.add_paragraph("Qualifications Not Provided", style='List Bullet')
        p.style.font.size = Pt(10)

    doc.add_heading("Skills", level=1)
    if resume_data["skills"]:
        for skill in resume_data["skills"]:
            p = doc.add_paragraph(skill, style='List Bullet')
            p.style.font.size = Pt(10)
    else:
        p = doc.add_paragraph("Skills Not Provided", style='List Bullet')
        p.style.font.size = Pt(10)

    if resume_data["positions"]:
        doc.add_heading("Extracurriculars", level=1)
        for pos in resume_data["positions"]:
            p = doc.add_paragraph(pos, style='List Bullet')
            p.style.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    if return_pdf:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(buffer.getvalue())
            tmp_path = tmp.name
        pdf_path = tmp_path.replace(".docx", ".pdf")
        docx2pdf.convert(tmp_path, pdf_path)
        with open(pdf_path, "rb") as f:
            pdf_buffer = BytesIO(f.read())
        os.remove(tmp_path)
        os.remove(pdf_path)
        return pdf_buffer
    return buffer

def generate_intermediate_word_resume(return_pdf=False):
    resume_data = st.session_state["resume_data"]
    doc = DocxDocument()

    name = doc.add_paragraph()
    run = name.add_run(resume_data["personal_info"]["full_name"] or "Full Name Not Provided\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(190, 130, 0)

    job_title = resume_data["experience"][0]["job_title"] if resume_data["experience"] else "Intermediate Professional Not Provided"
    title = doc.add_paragraph()
    title_run = title.add_run(job_title)
    title_run.font.size = Pt(14)
    title_run.bold = True

    contact = doc.add_paragraph()
    contact_parts = []
    if resume_data["personal_info"]["phone"]:
        contact_parts.append(f"Phone: {resume_data['personal_info']['phone']}")
    else:
        contact_parts.append("Phone: Not Provided")
    if resume_data["personal_info"]["email"]:
        contact_parts.append(f"Email: {resume_data['personal_info']['email']}")
    else:
        contact_parts.append("Email: Not Provided")
    if resume_data["personal_info"]["linkedin"]:
        contact_parts.append(f"LinkedIn: {resume_data['personal_info']['linkedin']}")
    if resume_data["personal_info"]["github"]:
        contact_parts.append(f"GitHub: {resume_data['personal_info']['github']}")
    contact.add_run(" | ".join(contact_parts) + "\n").font.size = Pt(10)
    contact.add_run(f"Address: {resume_data['personal_info']['address'] or 'Not Provided'}\n").font.size = Pt(10)

    summary_title = doc.add_paragraph()
    summary_title.add_run("SUMMARY").bold = True
    summary_title.runs[0].font.size = Pt(12)

    summary = doc.add_paragraph()
    summary.add_run(resume_data["summary"] or "Summary Not Provided").font.size = Pt(10)

    skills_title = doc.add_paragraph()
    skills_title.add_run("SKILLS").bold = True
    skills_title.runs[0].font.size = Pt(12)

    skills = doc.add_paragraph()
    skills_text = ", ".join(resume_data["skills"]) if resume_data["skills"] else "Skills Not Provided"
    skills.add_run(skills_text).font.size = Pt(10)

    work_title = doc.add_paragraph()
    work_title.add_run("WORK EXPERIENCE").bold = True
    work_title.runs[0].font.size = Pt(12)

    if resume_data["experience"]:
        for exp in resume_data["experience"][:2]:
            job = doc.add_paragraph()
            job.add_run(f"{exp.get('job_title', 'Title Not Provided')}\n").bold = True
            job.add_run(f"{exp.get('company', 'Company Not Provided')} – {exp.get('dates', 'Dates Not Provided')}\n").italic = True

            if exp.get("responsibilities"):
                for resp in exp["responsibilities"].split("; "):
                    doc.add_paragraph(resp, style='List Bullet').style.font.size = Pt(10)
            if exp.get("achievements"):
                for ach in exp["achievements"].split("; "):
                    doc.add_paragraph(ach, style='List Bullet').style.font.size = Pt(10)
    else:
        job = doc.add_paragraph()
        job.add_run("Intermediate Professional Not Provided\n").bold = True
        job.add_run("Company Not Provided – Dates Not Provided\n").italic = True
        doc.add_paragraph("Experience Not Provided", style='List Bullet').style.font.size = Pt(10)

    edu_title = doc.add_paragraph()
    edu_title.add_run("EDUCATION").bold = True
    edu_title.runs[0].font.size = Pt(12)

    education = doc.add_paragraph()
    edu_text = ", ".join(resume_data["qualifications"]) if resume_data["qualifications"] else "Education Not Provided"
    education.add_run(edu_text).font.size = Pt(10)

    if resume_data["projects"]:
        projects_title = doc.add_paragraph()
        projects_title.add_run("PROJECTS").bold = True
        projects_title.runs[0].font.size = Pt(12)
        projects = doc.add_paragraph()
        projects.add_run("\n".join(resume_data["projects"])).font.size = Pt(10)

    if resume_data["certifications"]:
        cert_title = doc.add_paragraph()
        cert_title.add_run("CERTIFICATIONS").bold = True
        cert_title.runs[0].font.size = Pt(12)
        certs = doc.add_paragraph()
        certs.add_run("\n".join(resume_data["certifications"])).font.size = Pt(10)

    if resume_data["positions"]:
        pos_title = doc.add_paragraph()
        pos_title.add_run("EXTRACURRICULARS").bold = True
        pos_title.runs[0].font.size = Pt(12)
        pos = doc.add_paragraph()
        pos.add_run("\n".join(resume_data["positions"])).font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    if return_pdf:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(buffer.getvalue())
            tmp_path = tmp.name
        pdf_path = tmp_path.replace(".docx", ".pdf")
        docx2pdf.convert(tmp_path, pdf_path)
        with open(pdf_path, "rb") as f:
            pdf_buffer = BytesIO(f.read())
        os.remove(tmp_path)
        os.remove(pdf_path)
        return pdf_buffer
    return buffer

def generate_veteran_pdf_resume():
    resume_data = st.session_state["resume_data"]
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    c.setFillColor(HexColor("#0077b6"))
    c.rect(0, height - 50, width, 50, fill=1, stroke=0)

    c.setFont("Helvetica-Bold", 18)
    c.setFillColor(black)
    c.drawString(50, height - 35, resume_data["personal_info"]["full_name"] or "Full Name Not Provided")
    c.setFont("Helvetica", 12)
    contact_parts = []
    if resume_data["personal_info"]["address"]:
        contact_parts.append(f"📍 {resume_data['personal_info']['address']}")
    else:
        contact_parts.append("📍 Address Not Provided")
    if resume_data["personal_info"]["phone"]:
        contact_parts.append(f"📞 {resume_data['personal_info']['phone']}")
    else:
        contact_parts.append("📞 Phone Not Provided")
    if resume_data["personal_info"]["email"]:
        contact_parts.append(f"✉ {resume_data['personal_info']['email']}")
    else:
        contact_parts.append("✉ Email Not Provided")
    if resume_data["personal_info"]["linkedin"]:
        contact_parts.append(f"🔗 {resume_data['personal_info']['linkedin']}")
    if resume_data["personal_info"]["github"]:
        contact_parts.append(f"🔗 {resume_data['personal_info']['github']}")
    contact_info = "  |  ".join(contact_parts)
    c.drawString(50, height - 55, contact_info)

    y = height - 100

    def add_section(title, content, y_pos):
        c.setFont("Helvetica-Bold", 14)
        c.setFillColor(HexColor("#0077b6"))
        c.drawString(50, y_pos, title)
        c.setFillColor(black)
        c.setFont("Helvetica", 10)
        y_pos -= 20
        if isinstance(content, list):
            for line in content:
                if line:  # Only include non-empty lines
                    c.drawString(50, y_pos, line)
                    y_pos -= 15
        else:
            c.drawString(50, y_pos, content)
            y_pos -= 15
        return y_pos - 20

    y = add_section("Summary", resume_data["summary"] or "Summary Not Provided", y)
    education_content = resume_data["qualifications"] if resume_data["qualifications"] else ["Education Not Provided"]
    y = add_section("Education", education_content, y)
    employment_content = [f"{exp.get('job_title', 'Title Not Provided')} - {exp.get('company', 'Company Not Provided')} ({exp.get('dates', 'Dates Not Provided')})" for exp in resume_data["experience"]] if resume_data["experience"] else ["Employment Not Provided"]
    y = add_section("Employment", employment_content, y)
    if resume_data["certifications"]:
        y = add_section("Certifications", resume_data["certifications"], y)
    achievements_content = [f"- {exp.get('achievements', '')}" for exp in resume_data["experience"] if exp.get('achievements')] if any(exp.get('achievements') for exp in resume_data["experience"]) else ["Achievements Not Provided"]
    y = add_section("Achievements", achievements_content, y)
    if resume_data["positions"]:
        y = add_section("Extracurriculars", resume_data["positions"], y)
    skills_content = ", ".join(resume_data["skills"]) if resume_data["skills"] else "Skills Not Provided"
    y = add_section("Skills", skills_content, y)
    languages_content = "English"
    y = add_section("Languages", languages_content, y)

    c.save()
    buffer.seek(0)
    return buffer

def generate_interview_questions():
    resume_data = st.session_state["resume_data"]
    resume_content = (
        f"Personal Info: {resume_data['personal_info']}\n"
        f"Summary: {resume_data['summary']}\n"
        f"Experience: {resume_data['experience']}\n"
        f"Projects: {resume_data['projects']}\n"
        f"Qualifications: {resume_data['qualifications']}\n"
        f"Skills: {resume_data['skills']}\n"
        f"Certifications: {resume_data['certifications']}\n"
        f"Positions of Responsibility: {resume_data['positions']}"
    )

    prompt = f"""
    Based on the following resume data, generate 20 specific interview questions tailored to the candidate's experience, skills, projects, education, certifications, and achievements. Ensure the questions are relevant, insightful, and encourage detailed responses about their background. Number each question from 1 to 20 and list them clearly.

    Resume Data:
    {resume_content}
    """

    response = get_gemini_response(prompt)
    if response:
        return response
    else:
        return "Failed to generate interview questions."

def preview_pdf_scrollable(pdf_buffer):
    pdf_buffer.seek(0)
    base64_pdf = base64.b64encode(pdf_buffer.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600px" type="application/pdf" title="Resume Preview"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def create_download_link(file_path_or_buffer, file_name, link_text):
    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "rb") as f:
            file_content = f.read()
    else:
        file_content = file_path_or_buffer.getvalue()
    b64 = base64.b64encode(file_content).decode()
    mime_type = "application/pdf" if file_name.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return f'<a href="data:{mime_type};base64,{b64}" download="{file_name}" target="_blank">{link_text}</a>'

st.markdown("""
    <style>
    .stTextInput, .stTextArea {
        margin-bottom: 20px;
        padding: 10px;
        width: 100%;
    }
    .stRadio > div {
        margin-bottom: 20px;
    }
    .stButton > button {
        margin-top: 10px;
        margin-bottom: 20px;
        padding: 10px 20px;
    }
    .stExpander {
        margin-bottom: 20px;
    }
    .question-section {
        font-size: 18px;
        line-height: 1.5;
        margin-bottom: 20px;
    }
    .language-instruction {
        font-size: 14px;
        color: #888;
        margin-bottom: 20px;
    }
    .personal-info-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 20px;
    }
    button:focus {
        outline: 2px solid #0078d4;
    }
    input:focus, textarea:focus {
        outline: 2px solid #0078d4;
    }
    </style>
""", unsafe_allow_html=True)

with st.container():
    col1, col2 = st.columns([1, 4])
    with col1:
        st.image("/Users/sachinbhat/Desktop/WhatsApp Image 2025-03-31 at 5.44.19 PM.jpeg", width=100, caption="ResWhisper Icon")
    with col2:
        st.title("ResuWhisper AI")
        st.write("Create a professional resume with your voice")

if st.session_state["page"] != "login" and st.session_state["page"] != "signup" and st.session_state["page"] != "welcome":
    progress_value = 0
    total_questions = len(questions)
    if st.session_state["page"] == "language_selection":
        progress_value = 10
    elif st.session_state["page"] == "consent":
        progress_value = 15
    elif st.session_state["page"] == "resume_template":
        progress_value = 20
    elif st.session_state["page"] == "questions":
        progress_value = 20 + (st.session_state["current_question_index"] / total_questions * 60)
    elif st.session_state["page"] == "preview":
        progress_value = 100
    st.progress(progress_value / 100, text=f"Progress: {int(progress_value)}%")

if st.session_state["page"] == "login":
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        user_id = login(username, password)
        if user_id:
            st.session_state["authenticated"] = True
            st.session_state["username"] = username
            session_id = start_new_session(user_id)
            if session_id:
                st.session_state["session_id"] = session_id
                st.session_state["page"] = "welcome"
                st.success("Logged in successfully!")
                st.rerun()
            else:
                st.error("Failed to start a new session")
        else:
            st.error("Invalid username or password")
    if st.button("Sign Up Instead"):
        st.session_state["page"] = "signup"
        st.rerun()

elif st.session_state["page"] == "signup":
    st.title("Sign Up")
    new_username = st.text_input("New Username")
    new_password = st.text_input("New Password", type="password")
    confirm_password = st.text_input("Confirm Password", type="password")
    if st.button("Sign Up"):
        if new_password != confirm_password:
            st.error("Passwords do not match")
        elif new_username and new_password:
            if signup(new_username, new_password):
                st.session_state["authenticated"] = True
                st.session_state["username"] = new_username
                user_id = login(new_username, new_password)
                if user_id:
                    session_id = start_new_session(user_id)
                    if session_id:
                        st.session_state["session_id"] = session_id
                        st.session_state["page"] = "welcome"
                        st.success("Signed up successfully!")
                        st.rerun()
                    else:
                        st.error("Failed to start a new session")
                else:
                    st.error("Failed to log in after signup")
            else:
                st.error("Signup failed: Username may already exist")
        else:
            st.error("Please fill all fields")
    if st.button("Back to Login"):
        st.session_state["page"] = "login"
        st.rerun()

elif st.session_state["page"] == "welcome" and st.session_state["authenticated"]:
    st.markdown("""
    ### Build Your Professional Resume with Just Your Voice!
    
    ResuWhisper AI uses advanced AI to help you create a professional resume simply by answering a few questions. 
    Speak naturally, and we'll convert your responses into a polished, well-formatted resume.
    
    **How it works:**
    1. Select a language for your responses
    2. Choose a resume template that fits your career stage
    3. Answer a series of questions using your voice or by uploading audio
    4. The AI processes your responses into professional resume content
    5. Preview and download your completed resume
    
    **Key Features:**
    - Voice-to-text conversion for easy input
    - AI-powered content formatting and enhancement
    - Multiple professional resume templates
    - Download in PDF or Word format
    """)
    
    if st.button("🚀 Get Started", key="welcome_start"):
        st.session_state["page"] = "language_selection"
        st.rerun()

elif st.session_state["page"] == "language_selection":
    st.title("🌍 Choose Your Preferred Language")
    st.write("Select the language you’ll use to answer the questions (will be transcribed to English):")
    
    cols = st.columns(4)
    for i, lang in enumerate(languages):
        with cols[i % 4]:
            if st.button(lang, key=f"lang_{lang}", help=f"Select {lang} as your language"):
                st.session_state["selected_language"] = lang
                if st.session_state["session_id"]:
                    update_session_data(st.session_state["session_id"], "selected_language", lang)
                if lang != "English":
                    translate_questions(lang)
                st.rerun()
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("← Back to Welcome", key="lang_back"):
            st.session_state["page"] = "welcome"
            st.rerun()
    with col2:
        if st.session_state["selected_language"]:
            if st.button("Next →", key="lang_next"):
                st.session_state["page"] = "consent"
                st.rerun()
        else:
            st.info("Please select a language")

elif st.session_state["page"] == "consent":
    st.title("📜 User Consent")
    st.markdown("""
    Before we proceed, please give us your consent to use your voice input for processing your resume.
    
    By checking the box below, you agree to allow this application to record and process your voice responses.
    """)
    st.session_state["consent_given"] = st.checkbox("I agree to the terms and conditions.", key="consent_checkbox")
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("← Back to Language Selection", key="consent_back"):
            st.session_state["page"] = "language_selection"
            st.rerun()
    with col2:
        if st.session_state["consent_given"]:
            if st.button("Proceed →", key="consent_next"):
                st.session_state["page"] = "resume_template"
                st.rerun()
        else:
            st.info("Please give your consent to proceed.")

elif st.session_state["page"] == "resume_template":
    st.title("📄 Choose Your Resume Template")
    st.write("Select a template that best matches your career stage:")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Entry Level / Student Template")
        st.image("https://jofibostorage.blob.core.windows.net/blog/Fresher-reverse.chronological.png", width=200, caption="Entry Level Template")
        st.write(resume_templates["Fresher"])
        if st.button("Select Entry Level Template", key="fresher_template"):
            st.session_state["resume_template"] = "Fresher"
            if st.session_state["session_id"]:
                update_session_data(st.session_state["session_id"], "resume_template", "Fresher")
            st.session_state["page"] = "questions"
            st.rerun()
        
        st.subheader("Executive / Senior Level Template")
        st.image("https://cdn-blog.novoresume.com/articles/executive-resume-example/executive-resume-sample.png", width=200, caption="Executive Template")
        st.write(resume_templates["Veteran"])
        if st.button("Select Executive Template", key="veteran_template"):
            st.session_state["resume_template"] = "Veteran"
            if st.session_state["session_id"]:
                update_session_data(st.session_state["session_id"], "resume_template", "Veteran")
            st.session_state["page"] = "questions"
            st.rerun()
    
    with col2:
        st.subheader("Mid-Career Professional Template")
        st.image("https://www.resumebuilder.com/wp-content/uploads/2020/12/Sales-Executive-Resume-Example-Banner-Image.png", width=200, caption="Mid-Career Template")
        st.write(resume_templates["Intermediate"])
        if st.button("Select Mid-Career Template", key="intermediate_template"):
            st.session_state["resume_template"] = "Intermediate"
            if st.session_state["session_id"]:
                update_session_data(st.session_state["session_id"], "resume_template", "Intermediate")
            st.session_state["page"] = "questions"
            st.rerun()
    
    if st.button("← Back to Consent", key="template_back"):
        st.session_state["page"] = "consent"
        st.rerun()

elif st.session_state["page"] == "questions":
    st.title("Resume Information Questionnaire")
    
    current_q = st.session_state["current_question_index"] + 1
    total_q = len(questions)
    st.write(f"Question {current_q} of {total_q}")
    
    st.header(section_headers[st.session_state["current_question_index"]])
    
    col1, col2 = st.columns([2, 3])

    with col1:
        current_index = st.session_state["current_question_index"]
        current_question = questions[current_index]
        st.markdown(f"### {current_question}")
        if st.session_state["selected_language"] != "English" and st.session_state["selected_language"] in st.session_state["translated_questions"]:
            translated_question = st.session_state["translated_questions"][st.session_state["selected_language"]][current_index]
            st.markdown(f"### {translated_question}")
        st.markdown(f'<div class="language-instruction">Please answer in {st.session_state["selected_language"]} (will be transcribed to English).</div>', unsafe_allow_html=True)

        input_option = st.radio("Choose Input Method:", ["Record Audio", "Upload Audio", "Text"], horizontal=True, key=f"input_method_{current_index}")

        if input_option == "Record Audio":
            if st.session_state["audio_file"] is None:
                st.session_state["audio_file"] = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name

            col_a, col_b = st.columns(2)
            with col_a:
                if not st.session_state.get("recording_state", False):
                    if st.button("🎤 Start Recording", key=f"start_rec_{current_index}"):
                        st.session_state["recording_state"] = True
                        stop_event = threading.Event()
                        st.session_state["stop_event"] = stop_event
                        recording_thread = threading.Thread(target=record_audio, args=(st.session_state["audio_file"], stop_event), daemon=True)
                        st.session_state["recording_thread"] = recording_thread
                        recording_thread.start()
                        st.rerun()
            with col_b:
                if st.session_state.get("recording_state", False):
                    st.warning("🔴 Recording in progress... speak clearly!")
                    if st.button("⏹ Stop Recording", key=f"stop_rec_{current_index}"):
                        st.session_state["recording_state"] = False
                        if "stop_event" in st.session_state:
                            st.session_state["stop_event"].set()
                        if "recording_thread" in st.session_state and st.session_state["recording_thread"].is_alive():
                            st.session_state["recording_thread"].join(timeout=5)
                        if os.path.exists(st.session_state["audio_file"]) and os.path.getsize(st.session_state["audio_file"]) > 0:
                            st.success(f"✅ Recording saved to {st.session_state['audio_file']}!")
                            with st.spinner("Transcribing..."):
                                progress = st.progress(0)
                                prompt = f"Transcribe the audio to English, assuming it is spoken in {st.session_state['selected_language']}."
                                response = get_gemini_response(prompt, st.session_state["audio_file"])
                                progress.progress(100)
                                if response:
                                    st.session_state["current_response"] = response
                                    update_resume_data(current_index, response)
                                else:
                                    st.error("Transcription failed after recording")
                        else:
                            st.error("Audio file was not created or is empty")
                        if "recording_thread" in st.session_state:
                            del st.session_state["recording_thread"]
                        if "stop_event" in st.session_state:
                            del st.session_state["stop_event"]
                        st.rerun()

        elif input_option == "Upload Audio":
            uploaded_file = st.file_uploader("Upload an audio file (MP3, WAV)", type=["mp3", "wav"], key=f"upload_{current_index}")
            if uploaded_file:
                if st.session_state["audio_file"] is None:
                    st.session_state["audio_file"] = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
                with open(st.session_state["audio_file"], "wb") as f:
                    f.write(uploaded_file.getvalue())
                if os.path.exists(st.session_state["audio_file"]) and os.path.getsize(st.session_state["audio_file"]) > 0:
                    st.success("✅ Audio file uploaded!")
                    if current_index not in st.session_state["transcribed_once"]:
                        with st.spinner("Transcribing..."):
                            progress = st.progress(0)
                            prompt = f"Transcribe the audio to English, assuming it is spoken in {st.session_state['selected_language']}."
                            response = get_gemini_response(prompt, st.session_state["audio_file"])
                            progress.progress(100)
                            if response:
                                st.session_state["current_response"] = response
                                update_resume_data(current_index, response)
                                st.session_state["transcribed_once"][current_index] = True
                            else:
                                st.error("Transcription failed")
                    else:
                        st.info("Audio already transcribed. Edit your response below if needed.")
                else:
                    st.error("Failed to save uploaded file")

        elif input_option == "Text":
            text_input = st.text_area("Type your response here:", height=200, key=f"text_{current_index}")
            if st.button("Submit Text", key=f"submit_text_{current_index}"):
                if text_input:
                    st.session_state["current_response"] = text_input
                    update_resume_data(current_index, text_input)
                else:
                    st.error("Please enter some text.")

        if st.session_state["current_response"]:
            st.subheader("Your Response:")
            edited_response = st.text_area("Edit your response if needed:", st.session_state["current_response"], height=200, key=f"trans_{current_index}")
            can_proceed = True
            if current_index not in [3, 6, 7]:
                if not edited_response.strip():
                    st.error("This question is mandatory. Please provide a response.")
                    can_proceed = False
            if st.button("✅ Process Response and Continue", key=f"process_{current_index}", disabled=not can_proceed):
                st.session_state["current_response"] = edited_response
                st.session_state["responses"][current_index] = edited_response
                if st.session_state["session_id"]:
                    update_session_data(st.session_state["session_id"], "responses", st.session_state["responses"])
                if current_index < len(questions) - 1:
                    st.session_state["current_question_index"] += 1
                    if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):
                        os.remove(st.session_state["audio_file"])
                    st.session_state["audio_file"] = None
                    st.session_state["current_response"] = None
                    st.rerun()
                else:
                    st.session_state["page"] = "preview"
                    st.rerun()

        col_nav1, col_nav2, col_nav3 = st.columns([1, 1, 1])
        with col_nav1:
            if st.session_state["current_question_index"] > 0:
                if st.button("← Previous Question", key=f"prev_{current_index}"):
                    st.session_state["current_question_index"] -= 1
                    st.session_state["current_response"] = None
                    if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):
                        os.remove(st.session_state["audio_file"])
                    st.session_state["audio_file"] = None
                    st.rerun()
        with col_nav3:
            if current_index in [3, 6, 7]:
                if st.button("Skip This Question →", key=f"skip_{current_index}"):
                    if current_index < len(questions) - 1:
                        st.session_state["current_question_index"] += 1
                        st.session_state["current_response"] = None
                        if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):
                            os.remove(st.session_state["audio_file"])
                        st.session_state["audio_file"] = None
                        st.rerun()
                    else:
                        st.session_state["page"] = "preview"
                        st.rerun()

    with col2:
        st.subheader("Live Resume Editor")
        st.markdown("**You can edit the Live Resume Editor if you are not satisfied. Type it or re-record yourself, it all works!**")
        resume_data = st.session_state["resume_data"]

        with st.expander("📋 Personal Information", expanded=True):
            st.markdown('<div class="personal-info-grid">', unsafe_allow_html=True)
            col_a, col_b = st.columns(2)
            with col_a:
                new_full_name = st.text_input("Full Name", resume_data["personal_info"]["full_name"], key="personal_full_name")
                resume_data["personal_info"]["full_name"] = new_full_name
            with col_b:
                new_degree = st.text_input("Degree", resume_data["personal_info"]["degree"], key="personal_degree")
                resume_data["personal_info"]["degree"] = new_degree
            col_c, col_d = st.columns(2)
            with col_c:
                new_phone = st.text_input("Phone", resume_data["personal_info"]["phone"], key="personal_phone")
                resume_data["personal_info"]["phone"] = new_phone
            with col_d:
                new_email = st.text_input("Email", resume_data["personal_info"]["email"], key="personal_email")
                resume_data["personal_info"]["email"] = new_email
            col_e, col_f = st.columns(2)
            with col_e:
                new_linkedin = st.text_input("LinkedIn (Optional)", resume_data["personal_info"]["linkedin"], key="personal_linkedin")
                resume_data["personal_info"]["linkedin"] = new_linkedin
            with col_f:
                new_github = st.text_input("GitHub (Optional)", resume_data["personal_info"]["github"], key="personal_github")
                resume_data["personal_info"]["github"] = new_github
            col_g, _ = st.columns(2)
            with col_g:
                new_address = st.text_input("Address", resume_data["personal_info"]["address"], key="personal_address")
                resume_data["personal_info"]["address"] = new_address
            st.markdown('</div>', unsafe_allow_html=True)

            if not all([resume_data["personal_info"]["full_name"], resume_data["personal_info"]["degree"],
                        resume_data["personal_info"]["phone"], resume_data["personal_info"]["email"],
                        resume_data["personal_info"]["address"]]):
                st.error("All fields except LinkedIn and GitHub are mandatory.")

        with st.expander("💼 Summary", expanded=True):
            new_summary = st.text_area("Summary", resume_data["summary"], height=100, key=f"summary_{current_index}")
            resume_data["summary"] = new_summary
            if current_index >= 1 and not new_summary.strip():
                st.error("Summary is mandatory.")

        if st.session_state["current_question_index"] >= 2:
            with st.expander("👔 Experience", expanded=True):
                if not resume_data["experience"]:
                    resume_data["experience"] = [{"job_title": "", "company": "", "dates": "", "responsibilities": "", "achievements": ""}]
                for i, exp in enumerate(resume_data["experience"]):
                    st.write(f"Experience {i+1}")
                    exp["job_title"] = st.text_input(f"Job Title {i+1}", exp.get("job_title", ""), key=f"exp_title_{i}")
                    exp["company"] = st.text_input(f"Company {i+1}", exp.get("company", ""), key=f"exp_company_{i}")
                    exp["dates"] = st.text_input(f"Dates {i+1}", exp.get("dates", ""), key=f"exp_dates_{i}")
                    exp["responsibilities"] = st.text_area(f"Responsibilities {i+1}", exp.get("responsibilities", ""), key=f"exp_resp_{i}")
                    exp["achievements"] = st.text_area(f"Achievements {i+1}", exp.get("achievements", ""), key=f"exp_achieve_{i}")
                if not any(exp["job_title"].strip() for exp in resume_data["experience"]):
                    st.error("Experience is mandatory.")

        if st.session_state["current_question_index"] >= 3:
            with st.expander("🚀 Projects (Optional)", expanded=True):
                if not resume_data["projects"]:
                    resume_data["projects"] = [""]
                for i, proj in enumerate(resume_data["projects"]):
                    new_proj = st.text_input(f"Project {i+1}", proj, key=f"proj_{i}")
                    resume_data["projects"][i] = new_proj

        if st.session_state["current_question_index"] >= 4:
            with st.expander("🎓 Qualifications", expanded=True):
                if not resume_data["qualifications"]:
                    resume_data["qualifications"] = [""]
                for i, qual in enumerate(resume_data["qualifications"]):
                    new_qual = st.text_input(f"Qualification {i+1}", qual, key=f"qual_{i}")
                    resume_data["qualifications"][i] = new_qual
                if not any(qual.strip() for qual in resume_data["qualifications"]):
                    st.error("Qualifications are mandatory.")

        if st.session_state["current_question_index"] >= 5:
            with st.expander("🛠️ Skills", expanded=True):
                if not resume_data["skills"]:
                    resume_data["skills"] = [""]
                for i, skill in enumerate(resume_data["skills"]):
                    new_skill = st.text_input(f"Skill {i+1}", skill, key=f"skill_{i}")
                    resume_data["skills"][i] = new_skill
                if not any(skill.strip() for skill in resume_data["skills"]):
                    st.error("Skills are mandatory.")

        if st.session_state["current_question_index"] >= 6:
            with st.expander("🏆 Certifications (Optional)", expanded=True):
                if not resume_data["certifications"]:
                    resume_data["certifications"] = [""]
                for i, cert in enumerate(resume_data["certifications"]):
                    new_cert = st.text_input(f"Certification {i+1}", cert, key=f"cert_{i}")
                    resume_data["certifications"][i] = new_cert

        if st.session_state["current_question_index"] >= 7:
            with st.expander("🌟 Positions of Responsibility (Optional)", expanded=True):
                if not resume_data["positions"]:
                    resume_data["positions"] = [""]
                for i, pos in enumerate(resume_data["positions"]):
                    new_pos = st.text_input(f"Position {i+1}", pos, key=f"pos_{i}")
                    resume_data["positions"][i] = new_pos

        st.session_state["resume_data"] = resume_data
        if st.session_state["session_id"]:
            update_session_data(st.session_state["session_id"], "resume_data", resume_data)

elif st.session_state["page"] == "preview":
    st.title("Your Resume Preview")
    
    st.subheader("Interview Questions Based on Your Resume")
    with st.spinner("Generating interview questions..."):
        interview_questions = generate_interview_questions()
        st.markdown(interview_questions)
        st.download_button("Download Questions", interview_questions, "interview_questions.txt", key="download_questions")

    st.subheader("PDF Preview")
    with st.spinner("Generating PDF Preview..."):
        if st.session_state["resume_template"] == "Fresher":
            pdf_buffer = generate_word_resume(return_pdf=True)
            preview_pdf_scrollable(pdf_buffer)
        elif st.session_state["resume_template"] == "Intermediate":
            pdf_buffer = generate_intermediate_word_resume(return_pdf=True)
            preview_pdf_scrollable(pdf_buffer)
        elif st.session_state["resume_template"] == "Veteran":
            pdf_buffer = generate_veteran_pdf_resume()
            preview_pdf_scrollable(pdf_buffer)
        else:
            st.error("Please select a resume template to preview.")

    st.subheader("Download Your Resume")
    file_format = st.selectbox("Choose format:", ["PDF", "Word"], key="download_format")
    
    if st.button("📥 Generate Resume for Download", key="download_button"):
        if st.session_state["resume_template"] == "Fresher":
            if file_format == "PDF":
                with st.spinner("Generating PDF..."):
                    pdf_buffer = generate_word_resume(return_pdf=True)
                    if st.session_state["session_id"]:
                        update_session_data(st.session_state["session_id"], "final_resume", pdf_buffer.getvalue())
                        update_session_data(st.session_state["session_id"], "final_resume_format", "PDF")
                    download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")
                    st.markdown(download_link, unsafe_allow_html=True)
            elif file_format == "Word":
                with st.spinner("Generating Word document..."):
                    word_buffer = generate_word_resume(return_pdf=False)
                    if st.session_state["session_id"]:
                        update_session_data(st.session_state["session_id"], "final_resume", word_buffer.getvalue())
                        update_session_data(st.session_state["session_id"], "final_resume_format", "Word")
                    download_link = create_download_link(word_buffer, "resume.docx", "Download Word Document")
                    st.markdown(download_link, unsafe_allow_html=True)
        elif st.session_state["resume_template"] == "Intermediate":
            if file_format == "PDF":
                with st.spinner("Generating PDF..."):
                    pdf_buffer = generate_intermediate_word_resume(return_pdf=True)
                    if st.session_state["session_id"]:
                        update_session_data(st.session_state["session_id"], "final_resume", pdf_buffer.getvalue())
                        update_session_data(st.session_state["session_id"], "final_resume_format", "PDF")
                    download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")
                    st.markdown(download_link, unsafe_allow_html=True)
            elif file_format == "Word":
                with st.spinner("Generating Word document..."):
                    word_buffer = generate_intermediate_word_resume(return_pdf=False)
                    if st.session_state["session_id"]:
                        update_session_data(st.session_state["session_id"], "final_resume", word_buffer.getvalue())
                        update_session_data(st.session_state["session_id"], "final_resume_format", "Word")
                    download_link = create_download_link(word_buffer, "resume.docx", "Download Word Document")
                    st.markdown(download_link, unsafe_allow_html=True)
        elif st.session_state["resume_template"] == "Veteran":
            with st.spinner("Generating PDF..."):
                pdf_buffer = generate_veteran_pdf_resume()
                if st.session_state["session_id"]:
                    update_session_data(st.session_state["session_id"], "final_resume", pdf_buffer.getvalue())
                    update_session_data(st.session_state["session_id"], "final_resume_format", "PDF")
                download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")
                st.markdown(download_link, unsafe_allow_html=True)
        else:
            st.error("Please select a resume template to download.")

    st.markdown("---")
    if st.button("← Back to Questionnaire", key="preview_back"):
        st.session_state["page"] = "questions"
        st.rerun()
    
    if st.button("🔄 Start Over", key="start_over"):
        for key in list(st.session_state.keys()):
            if key != "consent_given" and key != "authenticated" and key != "username" and key != "session_id":
                if key == "page":
                    st.session_state[key] = "welcome"
                elif key == "responses":
                    st.session_state[key] = {}
                elif key == "current_question_index":
                    st.session_state[key] = 0
                elif key == "resume_data":
                    st.session_state[key] = {
                        "personal_info": {
                            "full_name": "",
                            "degree": "",
                            "phone": "",
                            "email": "",
                            "linkedin": "",
                            "github": "",
                            "address": ""
                        },
                        "summary": "",
                        "qualifications": [],
                        "certifications": [],
                        "skills": [],
                        "experience": [],
                        "projects": [],
                        "positions": []
                    }
                elif key in ["resume_template", "current_response", "selected_language"]:
                    st.session_state[key] = None
        if st.session_state["session_id"]:
            update_session_data(st.session_state["session_id"], "responses", st.session_state["responses"])
            update_session_data(st.session_state["session_id"], "resume_data", st.session_state["resume_data"])
            update_session_data(st.session_state["session_id"], "resume_template", None)
            update_session_data(st.session_state["session_id"], "selected_language", None)
        st.rerun()

st.markdown("---")
st.markdown("""
<div style="text-align: center;">
    <p>© 2025 Resume Builder AI | Powered by CSSTUV</p>
</div>
""", unsafe_allow_html=True)