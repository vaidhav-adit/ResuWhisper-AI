import streamlit as st  # Imports the Streamlit library for creating web apps
import google.generativeai as genai  # Imports Google's generative AI library
import tempfile  # Imports module for creating temporary files
import sounddevice as sd  # Imports library for audio recording
import numpy as np  # Imports NumPy for numerical operations
import wave  # Imports module for handling WAV audio files
import threading  # Imports module for running tasks in parallel
import time  # Imports module for time-related functions
import os  # Imports module for interacting with the operating system
from pathlib import Path  # Imports Path class for file path handling
import base64  # Imports module for base64 encoding/decoding
from docx import Document as DocxDocument  # Imports python-docx library for Word documents
from docx.shared import Pt, Inches, RGBColor  # Imports units and color classes for Word docs
from docx.oxml.ns import qn  # Imports namespace utilities for Word docs
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Imports alignment options for Word docs
import docx2pdf  # Imports library to convert Word docs to PDF
from reportlab.lib.pagesizes import letter  # Imports letter page size for PDFs
from reportlab.pdfgen import canvas  # Imports canvas for PDF generation
from reportlab.lib.colors import black, HexColor  # Imports color options for PDFs
from io import BytesIO  # Imports BytesIO for in-memory file handling

# Replace with your actual API key
Google_API_Key = "AIzaSyCw8tLyHeobBO65GGnkLUVCGSMLdg-HsBw"  # Sets Google API key (replace with real key)
genai.configure(api_key=Google_API_Key)  # Configures the generative AI with the API key
model = genai.GenerativeModel('gemini-1.5-flash')  # Initializes the Gemini 1.5 Flash model

# Set page configuration for a professional look
st.set_page_config(  # Configures Streamlit page settings
    page_title="RESUWHISPER AI",  # Sets the page title
    page_icon="📝",  # Sets the page icon
    layout="wide",  # Sets the layout to wide
    initial_sidebar_state="expanded"  # Expands the sidebar by default
)

# Resume templates and other constants
resume_templates = {  # Defines a dictionary of resume templates
    "Fresher": "Ideal for 0-2 years of experience. Maximum 1 page.",  # Fresher template description
    "Intermediate": "Best for 3-7 years of experience. Maximum 1 page.",  # Intermediate template description
    "Veteran": "For 7+ years of experience. Maximum of 2 pages."  # Veteran template description
}

languages = [  # Lists supported languages for input
    "English", "Hindi", "Konkani", "Kannada", "Dogri", "Bodo", "Urdu", "Tamil",
    "Kashmiri", "Assamese", "Bengali", "Marathi", "Sindhi", "Maithili",
    "Punjabi", "Malayalam", "Manipuri", "Telugu", "Sanskrit", "Nepali",
    "Santali", "Gujarati", "Odia"
]

questions = [  # Defines list of questions for resume data collection
    "What is your full name, age, address, phone number, email, and LinkedIn or GitHub profile link (if any)?",
    "What are your career goals, key strengths, and professional personality traits?",
    "What is your work experience? Include job titles, company names, employment dates, responsibilities, and achievements.",
    "Tell us about the projects you have done. Include project names, descriptions, and your contributions (optional).",
    "What is your educational background? Include degrees, institutions, graduation years, and any relevant coursework or honors.",
    "What are your hard skills (e.g., technical skills) and soft skills (e.g., communication, teamwork)?",
    "What certifications do you have? Include the certification name, issuing organization, and date received.",
    "What are your extracurricular activities or recognitions? Include activities, organizations, dates, and achievements."
]

section_headers = [  # Defines headers for resume sections
    "📋 Personal Information",
    "💼 Professional Summary",
    "👔 Work Experience",
    "🚀 Projects",
    "🎓 Education",
    "🛠️ Skills",
    "🏆 Certifications",
    "🌟 Extracurricular Activities"
]

# Simulated database for users (replace with actual DB in production)
if "users_db" not in st.session_state:  # Checks if users_db exists in session state
    st.session_state["users_db"] = {}  # Initializes an empty dictionary for user data

# Initialize session states
def init_session_state():  # Defines function to initialize session state variables
    defaults = {  # Sets default values for session state
        "page": "login",  # Sets initial page to login
        "authenticated": False,  # Sets user as not authenticated
        "username": None,  # Sets username to None
        "selected_language": None,  # Sets selected language to None
        "consent_given": False,  # Sets consent flag to False
        "resume_template": None,  # Sets resume template to None
        "current_question_index": 0,  # Sets current question index to 0
        "responses": {},  # Initializes empty dictionary for responses
        "recording_state": False,  # Sets recording state to False
        "audio_file": None,  # Sets audio file path to None
        "stop_event": None,  # Sets stop event for recording to None
        "current_response": None,  # Sets current response to None
        "resume_data": {  # Initializes resume data structure
            "personal_info": {
                "full_name": "",
                "degree": "",
                "phone": "",
                "email": "",
                "linkedin": "",
                "github": "",
                "address": ""
            },
            "summary": "",
            "qualifications": [],
            "certifications": [],
            "skills": [],
            "experience": [],
            "projects": [],
            "positions": []
        },
        "transcribed_once": {},  # Tracks if audio has been transcribed
        "translated_questions": {}  # Stores translated questions
    }
    for key, value in defaults.items():  # Loops through defaults
        if key not in st.session_state:  # Checks if key is missing
            st.session_state[key] = value  # Sets default value

init_session_state()  # Calls function to initialize session state

def record_audio(filename, stop_event: threading.Event, samplerate=44100):  # Defines audio recording function
    audio_data = []  # Initializes list for audio data
    
    def callback(indata, frames, time, status):  # Defines callback for audio stream
        if not stop_event.is_set():  # Checks if recording should continue
            audio_data.append(indata.copy())  # Adds audio data to list
        else:
            raise sd.StopStream()  # Stops stream if event is set

    try:
        with sd.InputStream(samplerate=samplerate, channels=1, dtype=np.int16, callback=callback):  # Opens audio stream
            st.write("🎤 Recording... Press 'Stop Recording' to end.")  # Displays recording message
            while not stop_event.is_set():  # Loops until stop event
                time.sleep(0.1)  # Pauses briefly

        st.write("✅ Recording Completed.")  # Displays completion message
        audio_data = np.concatenate(audio_data, axis=0)  # Combines audio data into single array

        with wave.open(filename, 'wb') as wf:  # Opens WAV file for writing
            wf.setnchannels(1)  # Sets number of channels to 1
            wf.setsampwidth(2)  # Sets sample width to 2 bytes
            wf.setframerate(samplerate)  # Sets frame rate
            wf.writeframes(audio_data.tobytes())  # Writes audio data to file

        if os.path.exists(filename) and os.path.getsize(filename) > 0:  # Checks if file exists and has data
            st.success(f"✅ Recording saved to {filename}!")  # Displays success message
        else:
            st.error("Audio file was not created or is empty")  # Displays error if file is missing/empty
    except sd.StopStream:  # Handles stream stop exception
        pass
    except Exception as e:  # Catches other exceptions
        st.error(f"Recording error: {str(e)}")  # Displays error message

def get_gemini_response(input_msg, audio_path=None, mime_type="audio/wav"):  # Defines function to get AI response
    try:
        if audio_path:  # Checks if audio path is provided
            with open(audio_path, "rb") as f:  # Opens audio file in binary mode
                audio = genai.upload_file(audio_path, mime_type=mime_type)  # Uploads audio to Gemini
            response = model.generate_content([audio, input_msg])  # Generates response with audio and message
        else:
            response = model.generate_content([input_msg])  # Generates response with message only
        return response.text if response else None  # Returns response text or None
    except Exception as e:  # Catches exceptions
        st.error(f"Error with Gemini: {str(e)}")  # Displays error message
        return None  # Returns None on error

def translate_questions(language):  # Defines function to translate questions
    if language not in st.session_state["translated_questions"]:  # Checks if language is untranslated
        prompt = f"""  # Defines prompt for translation
        Translate the following 8 English questions into {language}. Provide the output as a numbered list, with each translation corresponding to the original question. Do not add extra commentary or modify the questions beyond translation.

        1. What is your full name, age, address, phone number, email, and LinkedIn or GitHub profile link (if any)?
        2. What are your career goals, key strengths, and professional personality traits?
        3. What is your work experience? Include job titles, company names, employment dates, responsibilities, and achievements.
        4. Tell us about the projects you have done. Include project names, descriptions, and your contributions (optional).
        5. What is your educational background? Include degrees, institutions, graduation years, and any relevant coursework or honors.
        6. What are your hard skills (e.g., technical skills) and soft skills (e.g., communication, teamwork)?
        7. What certifications do you have? Include the certification name, issuing organization, and date received.
        8. What are your extracurricular activities or recognitions? Include activities, organizations, dates, and achievements.
        """
        response = get_gemini_response(prompt)  # Gets translated response from Gemini
        if response:  # Checks if response exists
            translated = [line.split(". ", 1)[1] for line in response.strip().split("\n") if line.strip()]  # Extracts translations
            st.session_state["translated_questions"][language] = translated[:8]  # Stores up to 8 translations
        else:
            st.session_state["translated_questions"][language] = questions  # Falls back to English if translation fails
    return st.session_state["translated_questions"][language]  # Returns translated questions

def process_response_with_gemini(question_index, response):  # Defines function to process user responses with AI
    prompts = [  # Defines list of prompts for each question
        f"""  # Prompt for personal info
        **Your Goal**: Your goal is to meticulously extract and enhance personal details from the user's response to create a flawless, professional resume entry, ensuring every detail is accurate and perfectly formatted for the live resume editor.
        **Your Role**: You are a professional resume maker with years of experience crafting polished, error-free resumes for clients across industries. Make sure you be as accurate as possible.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to extract only what the user explicitly states—no assumptions, no inventions—and enhance it with professional formatting.
        **Detailed Instructions**:
        - Extract the following fields only if explicitly mentioned in the response:
          - Full Name (e.g., "John Doe")
          - Age (e.g., "22")
          - Address (e.g., "123 Main St, Springfield, IL, USA")
          - Phone Number (e.g., "+1-123-456-7890")
          - Email Address (e.g., "john.doe@example.com")
          - LinkedIn Profile (e.g., "linkedin.com/in/johndoe")
          - GitHub Profile (e.g., "github.com/johndoe")
        - Do not guess or add information the user didn’t provide. If a field is missing, leave it blank (e.g., "Age: ").
        - Enhance formatting for professionalism:
          - Standardize phone numbers (e.g., convert "1234567890" to "+1-123-456-7890").
          - Ensure email addresses are lowercase and valid (e.g., "John@EXAMPLE.com" becomes "john@example.com").
          - Verify URLs are complete (e.g., add "https://" to "linkedin.com/in/johndoe" if missing).
        - Avoid duplicating or misplacing fields (e.g., don’t put email in the address field).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        - Put what sounds like an "email" in the email field. What sounds like an address must be put int the address field. 
        **Example**:
        - User Input: "My name is John Doe, I live at 123 Main St, email is john@example, phone 123-456-7890, linkedin.com/in/johndoe"
        - Output:
          Full Name: John Doe
          Age: 
          Address: 123 Main St
          Phone Number: +1-123-456-7890
          Email Address: john@example.com
          LinkedIn Profile: https://linkedin.com/in/johndoe
          GitHub Profile: 
        **Output Format**: Return the extracted and enhanced data in this exact format, one field per line:
        Full Name: [Name]
        Age: [Age]
        Address: [Address]
        Phone Number: [Phone]
        Email Address: [Email]
        LinkedIn Profile: [LinkedIn]
        GitHub Profile: [GitHub]
        """,
        f"""  # Prompt for summary
        **Your Goal**: Your goal is to craft a concise, impactful, and professional summary for the user's resume based solely on their response, enhancing it to sound polished and tailored for the live resume editor without adding unmentioned details.
        **Your Role**: You are a professional resume maker skilled at transforming raw career aspirations and traits into compelling summaries that impress hiring managers.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to enhance only what the user explicitly mentions into a 50-70 word summary.
        **Detailed Instructions**:
        - Focus on three key areas if mentioned:
          - Career Goals (e.g., "Aspiring software developer")
          - Key Strengths (e.g., "problem-solving, coding")
          - Professional Traits (e.g., "team player, detail-oriented")
        - Do not invent details or add fluff. If the user doesn’t mention a category, omit it.
        - Enhance the language to be professional and concise:
          - Convert casual phrases (e.g., "I like coding") to polished ones (e.g., "Skilled in coding").
          - Combine elements logically (e.g., "I want to be a developer and I’m good at teamwork" becomes "Aspiring developer with strong teamwork skills").
        - Keep the summary between 50-70 words for brevity and impact.
        - Do not include suggestions or additional commentary—output only the enhanced summary.
        **Example**:
        - User Input: "I want to be a software engineer, I’m good at problem-solving and I work well in teams."
        - Output:
          Enhanced Summary: Aspiring software engineer with a strong aptitude for problem-solving and a proven ability to collaborate effectively in team environments. Passionate about leveraging technical skills to develop innovative solutions and contribute to organizational success.
        **Output Format**: Return the enhanced summary in this exact format:
        Enhanced Summary: [Summary]
        """,
        f"""  # Prompt for work experience
        **Your Goal**: Your goal is to extract and enhance the user’s work experience into detailed, structured fields for the live resume editor, ensuring each job is accurately separated and professionally formatted without errors or omissions.
        **Your Role**: You are a professional resume maker with expertise in organizing work history into clear, impactful entries that highlight achievements and responsibilities.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to break down multiple jobs if mentioned and enhance only what’s provided.
        **Detailed Instructions**:
        - For each job mentioned, extract these fields if explicitly stated:
          - Job Title (e.g., "Software Intern")
          - Company Name (e.g., "Tech Corp")
          - Dates (e.g., "June 2022 - August 2022")
          - Responsibilities (e.g., "Developed features using Python")
          - Achievements (e.g., "Improved performance by 20%")
        - Identify multiple jobs by keywords like "and," "then," "also," or distinct company names. Separate each job with "---".
        - Do not assume details not provided. Leave fields blank if missing (e.g., "Achievements: ").
        - Enhance professionalism:
          - Use action verbs (e.g., "Coded" becomes "Developed").
          - Quantify where possible if numbers are explicitly stated (e.g., "Improved performance by 20%").
        - Ensure no overlap or misplacement (e.g., dates don’t go into responsibilities).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I was a Software Intern at Tech Corp from June 2022 to August 2022, coded stuff, and then a Developer at XYZ Inc from Jan 2023 to now, built apps."
        - Output:
          Job Title: Software Intern
          Company Name: Tech Corp
          Dates: June 2022 - August 2022
          Responsibilities: Developed software features
          Achievements: 
          ---
          Job Title: Developer
          Company Name: XYZ Inc
          Dates: Jan 2023 - Present
          Responsibilities: Developed mobile applications
          Achievements: 
        **Output Format**: Return each job in this exact format, separated by "---":
        Job Title: [Title]
        Company Name: [Company]
        Dates: [Dates]
        Responsibilities: [Responsibilities]
        Achievements: [Achievements]
        ---
        """,
        f"""  # Prompt for projects
        **Your Goal**: Your goal is to extract and enhance the user’s project details into concise, professional entries for the live resume editor, treating this as an optional section and ensuring accuracy without adding unmentioned projects.
        **Your Role**: You are a professional resume maker adept at showcasing projects to highlight technical skills and contributions in a clear, employer-friendly format.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each project separately if multiple are mentioned.
        **Detailed Instructions**:
        - For each project mentioned, extract:
          - Project Name (e.g., "Portfolio Website")
          - Description (e.g., "Built a responsive site")
          - Contribution (e.g., "Coded frontend in React")
        - Identify multiple projects by keywords like "and," "also," or distinct names. List each with a "-".
        - Do not add projects not mentioned. If none are provided, return an empty response.
        - Enhance for clarity and professionalism:
          - Combine details into a single line (e.g., "Made a site" becomes "Portfolio Website: Built a responsive site, coded frontend").
          - Use technical terms where applicable (e.g., "Helped on app" becomes "Mobile App: Assisted in development, contributed to UI").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I did a Portfolio Website, built a site, and a Chat App, coded it with friends."
        - Output:
          - Portfolio Website: Built a responsive website, Coded frontend
          - Chat App: Developed a real-time application, Collaborated on coding
        **Output Format**: Return each project in this exact format, one per line:
        - [Project Name]: [Description], [Contribution]
        """,
        f"""  # Prompt for education
        **Your Goal**: Your goal is to extract and enhance the user’s educational background into detailed, accurate entries for the live resume editor, capturing every institution mentioned and formatting them professionally without missing or misplacing any details.
        **Your Role**: You are a professional resume maker with a keen eye for organizing educational history into a clear, chronological format that appeals to employers.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each educational entry separately.
        **Detailed Instructions**:
        - For each educational institution mentioned, extract:
          - Degree (e.g., "B.Tech in Computer Science")
          - Institution (e.g., "XYZ University")
          - Graduation Year (e.g., "2023")
          - Honors (e.g., "Cum Laude", if mentioned)
        - Identify multiple entries by keywords like "and," "also," "then," or distinct institutions. List each with a "-".
        - Do not add unmentioned details or fields (e.g., no "Relevant Coursework" unless stated).
        - Enhance for professionalism:
          - Standardize degree names (e.g., "CS degree" becomes "B.Tech in Computer Science").
          - Ensure years are four digits (e.g., "23" becomes "2023").
        - Order entries chronologically if dates are provided, most recent first.
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "B.Tech from IIT in 2023, Diploma from XYZ in 2020, and high school from ABC in 2018."
        - Output:
          - B.Tech in Computer Science, IIT, 2023
          - Diploma, XYZ, 2020
          - High School Diploma, ABC, 2018
        **Output Format**: Return each education entry in this exact format, one per line:
        - [Degree], [Institution], [Graduation Year], [Honors if mentioned]
        """,
        f"""  # Prompt for skills
        **Your Goal**: Your goal is to extract and enhance the user’s skills into a clear, categorized list for the live resume editor, ensuring every skill is accurately placed and professionally presented without errors or omissions.
        **Your Role**: You are a professional resume maker skilled at identifying and classifying skills to maximize their impact on a resume.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to separate hard and soft skills explicitly.
        **Detailed Instructions**:
        - Categorize skills into two groups if mentioned:
          - Hard Skills (technical, e.g., "Python, SQL, Machine Learning")
          - Soft Skills (interpersonal, e.g., "Communication, Teamwork")
        - Extract only skills the user states. Do not add unmentioned skills.
        - Enhance for clarity:
          - List skills as a comma-separated string (e.g., "coding, teamwork" becomes "Hard Skills: Python, Soft Skills: Teamwork").
          - Specify languages/tools for vague terms where explicitly stated (e.g., "Coding" becomes "Python" if Python is mentioned).
        - Avoid duplication or misclassification (e.g., "Teamwork" stays in soft skills, not hard).
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "I know Python, teamwork, and problem-solving."
        - Output:
          - Hard Skills: Python
          - Soft Skills: Teamwork, Problem-Solving
        **Output Format**: Return skills in this exact format, one category per line:
        - Hard Skills: [Hard Skills]
        - Soft Skills: [Soft Skills]
        """,
        f"""  # Prompt for certifications
        **Your Goal**: Your goal is to extract and enhance the user’s certifications into precise, professional entries for the live resume editor, ensuring every certification is captured accurately and formatted consistently without errors.
        **Your Role**: You are a professional resume maker experienced in presenting certifications to boost a candidate’s credibility.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each certification separately.
        **Detailed Instructions**:
        - For each certification mentioned, extract:
          - Certification Name (e.g., "AWS Certified Developer")
          - Issuing Organization (e.g., "Amazon")
          - Date (e.g., "July 2023")
        - Identify multiple certifications by keywords like "and," "also," or distinct names. List each with a "-".
        - Do not add unmentioned certifications. If none are provided, return an empty response.
        - Enhance for professionalism:
          - Standardize names (e.g., "AWS cert" becomes "AWS Certified Developer").
          - Format dates as "Month Year" (e.g., "7/23" becomes "July 2023").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "AWS cert from Amazon in July 2023, and Python cert from Coursera in 2022."
        - Output:
          - AWS Certified Developer, Amazon, July 2023
          - Python Certification, Coursera, 2022
        **Output Format**: Return each certification in this exact format, one per line:
        - [Certification Name], [Issuing Organization], [Date]
        """,
        f"""  # Prompt for extracurriculars
        **Your Goal**: Your goal is to extract and enhance the user’s extracurricular activities into detailed, professional entries for the live resume editor, treating this as optional and ensuring accuracy without adding unmentioned details.
        **Your Role**: You are a professional resume maker skilled at highlighting extracurriculars to showcase leadership and soft skills.
        **Context**: The user is building a resume and has provided the following response to the question: "{questions[question_index]}". The user's response is: "{response}". Your task is to list each activity separately if multiple are mentioned.
        **Detailed Instructions**:
        - For each activity mentioned, extract:
          - Title (e.g., "President of Coding Club")
          - Organization (e.g., "XYZ University")
          - Dates (e.g., "2022-2023")
          - Accomplishment (e.g., "Organized hackathon")
        - Identify multiple activities by keywords like "and," "also," or distinct titles. List each with a "-".
        - Do not add unmentioned activities. If none are provided, return an empty response.
        - Enhance for professionalism:
          - Use active language (e.g., "Helped club" becomes "Organized club events").
          - Standardize dates (e.g., "22-23" becomes "2022-2023").
        - Do not include suggestions or additional commentary—output only the enhanced data.
        **Example**:
        - User Input: "President of Coding Club at XYZ University 2022-2023, organized hackathon, and volunteer at ABC Org 2021."
        - Output:
          - President of Coding Club, XYZ University, 2022-2023, Organized hackathon
          - Volunteer, ABC Org, 2021
        **Output Format**: Return each activity in this exact format, one per line:
        - [Title], [Organization], [Dates], [Accomplishment]
        """
    ]

    prompt = prompts[question_index]  # Selects prompt based on question index
    result = get_gemini_response(prompt)  # Gets response from Gemini
    return result  # Returns processed result

def update_resume_data(question_index, response):  # Defines function to update resume data
    result = process_response_with_gemini(question_index, response)  # Processes response with Gemini
    if not result:  # Checks if result is empty
        st.error("Failed to process response with Gemini")  # Displays error
        return

    if question_index == 0:  # Handles personal info
        lines = result.split("\n")  # Splits result into lines
        personal_info = st.session_state["resume_data"]["personal_info"]  # Gets personal info dict
        for line in lines:  # Loops through lines
            if "Full Name:" in line:  # Checks for full name
                personal_info["full_name"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates full name
            elif "Address:" in line:  # Checks for address
                personal_info["address"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates address
            elif "Phone Number:" in line:  # Checks for phone
                personal_info["phone"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates phone
            elif "Email Address:" in line:  # Checks for email
                personal_info["email"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates email
            elif "LinkedIn Profile:" in line:  # Checks for LinkedIn
                personal_info["linkedin"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates LinkedIn
            elif "GitHub Profile:" in line:  # Checks for GitHub
                personal_info["github"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates GitHub
        st.session_state["resume_data"]["personal_info"] = personal_info  # Updates session state

    elif question_index == 1:  # Handles summary
        for line in result.split("\n"):  # Loops through lines
            if "Enhanced Summary:" in line:  # Checks for summary
                summary = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Extracts summary
                st.session_state["resume_data"]["summary"] = summary  # Updates summary
                break

    elif question_index == 2:  # Handles experience
        experience = []  # Initializes experience list
        current_exp = {}  # Initializes current experience dict
        for line in result.split("\n"):  # Loops through lines
            if line.strip() == "---" and current_exp:  # Checks for job separator
                experience.append(current_exp)  # Adds current job to list
                current_exp = {}  # Resets current job
            elif "Job Title:" in line:  # Checks for job title
                current_exp["job_title"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates job title
            elif "Company Name:" in line:  # Checks for company
                current_exp["company"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates company
            elif "Dates:" in line:  # Checks for dates
                current_exp["dates"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates dates
            elif "Responsibilities:" in line:  # Checks for responsibilities
                current_exp["responsibilities"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates responsibilities
            elif "Achievements:" in line:  # Checks for achievements
                current_exp["achievements"] = line.split(": ")[1] if len(line.split(": ")) > 1 else ""  # Updates achievements
        if current_exp:  # Checks if last job exists
            experience.append(current_exp)  # Adds last job to list
        st.session_state["resume_data"]["experience"] = experience  # Updates experience

    elif question_index == 3:  # Handles projects
        projects = []  # Initializes projects list
        for line in result.split("\n"):  # Loops through lines
            if line.startswith("-"):  # Checks for project line
                projects.append(line.strip("- ").strip())  # Adds project to list
        st.session_state["resume_data"]["projects"] = projects  # Updates projects

    elif question_index == 4:  # Handles qualifications
        qualifications = []  # Initializes qualifications list
        for line in result.split("\n"):  # Loops through lines
            if line.startswith("-"):  # Checks for qualification line
                qualifications.append(line.strip("- ").strip())  # Adds qualification to list
        st.session_state["resume_data"]["qualifications"] = qualifications  # Updates qualifications

    elif question_index == 5:  # Handles skills
        skills = []  # Initializes skills list
        for line in result.split("\n"):  # Loops through lines
            if line.startswith("-"):  # Checks for skill line
                skills.append(line.strip("- ").strip())  # Adds skill to list
        st.session_state["resume_data"]["skills"] = skills  # Updates skills

    elif question_index == 6:  # Handles certifications
        certifications = []  # Initializes certifications list
        for line in result.split("\n"):  # Loops through lines
            if line.startswith("-"):  # Checks for certification line
                certifications.append(line.strip("- ").strip())  # Adds certification to list
        st.session_state["resume_data"]["certifications"] = certifications  # Updates certifications

    elif question_index == 7:  # Handles extracurriculars
        positions = []  # Initializes positions list
        for line in result.split("\n"):  # Loops through lines
            if line.startswith("-"):  # Checks for position line
                positions.append(line.strip("- ").strip())  # Adds position to list
        st.session_state["resume_data"]["positions"] = positions  # Updates positions

def generate_word_resume(return_pdf=False):  # Defines function to generate Word resume
    resume_data = st.session_state["resume_data"]  # Gets resume data from session state
    doc = DocxDocument()  # Creates new Word document

    sections = doc.sections  # Gets document sections
    for section in sections:  # Loops through sections
        section.top_margin = Inches(0.3)  # Sets top margin
        section.bottom_margin = Inches(0.3)  # Sets bottom margin
        section.left_margin = Inches(0.3)  # Sets left margin
        section.right_margin = Inches(0.3)  # Sets right margin

    header = doc.add_paragraph()  # Adds header paragraph
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centers header
    run = header.add_run(resume_data["personal_info"]["full_name"] or "Full Name Not Provided")  # Adds name
    run.bold = True  # Makes name bold
    run.font.size = Pt(16)  # Sets name font size
    
    if resume_data["personal_info"]["degree"]:  # Checks if degree exists
        run = header.add_run(f"\n{resume_data['personal_info']['degree']}")  # Adds degree
    else:
        run = header.add_run("\nDegree Not Provided")  # Adds placeholder if no degree
    run.bold = True  # Makes degree bold
    run.font.size = Pt(12)  # Sets degree font size
    
    contact_info_parts = []  # Initializes contact info list
    if resume_data["personal_info"]["phone"]:  # Checks if phone exists
        contact_info_parts.append(resume_data["personal_info"]["phone"])  # Adds phone
    else:
        contact_info_parts.append("Phone Not Provided")  # Adds placeholder if no phone
    if resume_data["personal_info"]["email"]:  # Checks if email exists
        contact_info_parts.append(resume_data["personal_info"]["email"])  # Adds email
    else:
        contact_info_parts.append("Email Not Provided")  # Adds placeholder if no email
    if resume_data["personal_info"]["linkedin"]:  # Checks if LinkedIn exists
        contact_info_parts.append(resume_data["personal_info"]["linkedin"])  # Adds LinkedIn
    if resume_data["personal_info"]["github"]:  # Checks if GitHub exists
        contact_info_parts.append(resume_data["personal_info"]["github"])  # Adds GitHub
    contact_info = " | ".join(contact_info_parts)  # Joins contact info with separator
    run = header.add_run(f"\n{contact_info}")  # Adds contact info
    run.font.size = Pt(10)  # Sets contact info font size
    
    run = header.add_run(f"\n{resume_data['personal_info']['address'] or 'Address Not Provided'}")  # Adds address
    run.font.size = Pt(10)  # Sets address font size

    doc.add_heading("Summary", level=1)  # Adds Summary heading
    p = doc.add_paragraph(resume_data["summary"] or "Summary Not Provided")  # Adds summary text
    p.style.font.size = Pt(10)  # Sets summary font size

    doc.add_heading("Experience", level=1)  # Adds Experience heading
    if resume_data["experience"]:  # Checks if experience exists
        for exp in resume_data["experience"]:  # Loops through experience entries
            p = doc.add_paragraph(f"{exp.get('job_title', 'Title Not Provided')} at {exp.get('company', 'Company Not Provided')} ({exp.get('dates', 'Dates Not Provided')})")  # Adds job details
            p.style.font.size = Pt(10)  # Sets job details font size
            p.runs[0].bold = True  # Makes job title bold
            if exp.get("responsibilities"):  # Checks if responsibilities exist
                p = doc.add_paragraph(exp["responsibilities"], style='List Bullet')  # Adds responsibilities as bullets
                p.style.font.size = Pt(10)  # Sets responsibilities font size
            if exp.get("achievements"):  # Checks if achievements exist
                p = doc.add_paragraph(exp["achievements"], style='List Bullet')  # Adds achievements as bullets
                p.style.font.size = Pt(10)  # Sets achievements font size
    else:
        p = doc.add_paragraph("Experience Not Provided")  # Adds placeholder if no experience
        p.style.font.size = Pt(10)  # Sets placeholder font size

    if resume_data["projects"]:  # Checks if projects exist
        doc.add_heading("Projects", level=1)  # Adds Projects heading
        for proj in resume_data["projects"]:  # Loops through projects
            p = doc.add_paragraph(proj, style='List Bullet')  # Adds project as bullet
            p.style.font.size = Pt(10)  # Sets project font size

    doc.add_heading("Qualifications & Certifications", level=1)  # Adds Qualifications & Certifications heading
    all_qualifications = resume_data["qualifications"] + resume_data["certifications"]  # Combines qualifications and certifications
    if all_qualifications:  # Checks if qualifications exist
        for qual in all_qualifications:  # Loops through qualifications
            if qual:  # Checks if qualification is non-empty
                p = doc.add_paragraph(qual, style='List Bullet')  # Adds qualification as bullet
                p.style.font.size = Pt(10)  # Sets qualification font size
    else:
        p = doc.add_paragraph("Qualifications Not Provided", style='List Bullet')  # Adds placeholder if no qualifications
        p.style.font.size = Pt(10)  # Sets placeholder font size

    doc.add_heading("Skills", level=1)  # Adds Skills heading
    if resume_data["skills"]:  # Checks if skills exist
        for skill in resume_data["skills"]:  # Loops through skills
            p = doc.add_paragraph(skill, style='List Bullet')  # Adds skill as bullet
            p.style.font.size = Pt(10)  # Sets skill font size
    else:
        p = doc.add_paragraph("Skills Not Provided", style='List Bullet')  # Adds placeholder if no skills
        p.style.font.size = Pt(10)  # Sets placeholder font size

    if resume_data["positions"]:  # Checks if positions exist
        doc.add_heading("Extracurriculars", level=1)  # Adds Extracurriculars heading
        for pos in resume_data["positions"]:  # Loops through positions
            p = doc.add_paragraph(pos, style='List Bullet')  # Adds position as bullet
            p.style.font.size = Pt(10)  # Sets position font size

    buffer = BytesIO()  # Creates in-memory buffer
    doc.save(buffer)  # Saves document to buffer
    buffer.seek(0)  # Resets buffer position to start
    if return_pdf:  # Checks if PDF output is requested
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:  # Creates temp Word file
            tmp.write(buffer.getvalue())  # Writes buffer to temp file
            tmp_path = tmp.name  # Gets temp file path
        pdf_path = tmp_path.replace(".docx", ".pdf")  # Creates PDF path
        docx2pdf.convert(tmp_path, pdf_path)  # Converts Word to PDF
        with open(pdf_path, "rb") as f:  # Opens PDF file
            pdf_buffer = BytesIO(f.read())  # Reads PDF into buffer
        os.remove(tmp_path)  # Deletes temp Word file
        os.remove(pdf_path)  # Deletes temp PDF file
        return pdf_buffer  # Returns PDF buffer
    return buffer  # Returns Word buffer

def generate_intermediate_word_resume(return_pdf=False):  # Defines function for intermediate Word resume
    resume_data = st.session_state["resume_data"]  # Gets resume data
    doc = DocxDocument()  # Creates new Word document

    name = doc.add_paragraph()  # Adds name paragraph
    run = name.add_run(resume_data["personal_info"]["full_name"] or "Full Name Not Provided\n")  # Adds name
    run.bold = True  # Makes name bold
    run.font.size = Pt(22)  # Sets name font size
    run.font.color.rgb = RGBColor(190, 130, 0)  # Sets name color

    job_title = resume_data["experience"][0]["job_title"] if resume_data["experience"] else "Intermediate Professional Not Provided"  # Gets job title or placeholder
    title = doc.add_paragraph()  # Adds title paragraph
    title_run = title.add_run(job_title)  # Adds job title
    title_run.font.size = Pt(14)  # Sets title font size
    title_run.bold = True  # Makes title bold

    contact = doc.add_paragraph()  # Adds contact paragraph
    contact_parts = []  # Initializes contact parts list
    if resume_data["personal_info"]["phone"]:  # Checks if phone exists
        contact_parts.append(f"Phone: {resume_data['personal_info']['phone']}")  # Adds phone
    else:
        contact_parts.append("Phone: Not Provided")  # Adds placeholder if no phone
    if resume_data["personal_info"]["email"]:  # Checks if email exists
        contact_parts.append(f"Email: {resume_data['personal_info']['email']}")  # Adds email
    else:
        contact_parts.append("Email: Not Provided")  # Adds placeholder if no email
    if resume_data["personal_info"]["linkedin"]:  # Checks if LinkedIn exists
        contact_parts.append(f"LinkedIn: {resume_data['personal_info']['linkedin']}")  # Adds LinkedIn
    if resume_data["personal_info"]["github"]:  # Checks if GitHub exists
        contact_parts.append(f"GitHub: {resume_data['personal_info']['github']}")  # Adds GitHub
    contact.add_run(" | ".join(contact_parts) + "\n").font.size = Pt(10)  # Adds contact info
    contact.add_run(f"Address: {resume_data['personal_info']['address'] or 'Not Provided'}\n").font.size = Pt(10)  # Adds address

    summary_title = doc.add_paragraph()  # Adds summary title paragraph
    summary_title.add_run("SUMMARY").bold = True  # Adds SUMMARY text
    summary_title.runs[0].font.size = Pt(12)  # Sets summary title font size

    summary = doc.add_paragraph()  # Adds summary paragraph
    summary.add_run(resume_data["summary"] or "Summary Not Provided").font.size = Pt(10)  # Adds summary text

    skills_title = doc.add_paragraph()  # Adds skills title paragraph
    skills_title.add_run("SKILLS").bold = True  # Adds SKILLS text
    skills_title.runs[0].font.size = Pt(12)  # Sets skills title font size

    skills = doc.add_paragraph()  # Adds skills paragraph
    skills_text = ", ".join(resume_data["skills"]) if resume_data["skills"] else "Skills Not Provided"  # Joins skills or adds placeholder
    skills.add_run(skills_text).font.size = Pt(10)  # Adds skills text

    work_title = doc.add_paragraph()  # Adds work title paragraph
    work_title.add_run("WORK EXPERIENCE").bold = True  # Adds WORK EXPERIENCE text
    work_title.runs[0].font.size = Pt(12)  # Sets work title font size

    if resume_data["experience"]:  # Checks if experience exists
        for exp in resume_data["experience"][:2]:  # Loops through first two experiences
            job = doc.add_paragraph()  # Adds job paragraph
            job.add_run(f"{exp.get('job_title', 'Title Not Provided')}\n").bold = True  # Adds job title
            job.add_run(f"{exp.get('company', 'Company Not Provided')} – {exp.get('dates', 'Dates Not Provided')}\n").italic = True  # Adds company and dates

            if exp.get("responsibilities"):  # Checks if responsibilities exist
                for resp in exp["responsibilities"].split("; "):  # Splits responsibilities
                    doc.add_paragraph(resp, style='List Bullet').style.font.size = Pt(10)  # Adds responsibility as bullet
            if exp.get("achievements"):  # Checks if achievements exist
                for ach in exp["achievements"].split("; "):  # Splits achievements
                    doc.add_paragraph(ach, style='List Bullet').style.font.size = Pt(10)  # Adds achievement as bullet
    else:
        job = doc.add_paragraph()  # Adds job paragraph
        job.add_run("Intermediate Professional Not Provided\n").bold = True  # Adds placeholder title
        job.add_run("Company Not Provided – Dates Not Provided\n").italic = True  # Adds placeholder company/dates
        doc.add_paragraph("Experience Not Provided", style='List Bullet').style.font.size = Pt(10)  # Adds placeholder bullet

    edu_title = doc.add_paragraph()  # Adds education title paragraph
    edu_title.add_run("EDUCATION").bold = True  # Adds EDUCATION text
    edu_title.runs[0].font.size = Pt(12)  # Sets education title font size

    education = doc.add_paragraph()  # Adds education paragraph
    edu_text = ", ".join(resume_data["qualifications"]) if resume_data["qualifications"] else "Education Not Provided"  # Joins qualifications or adds placeholder
    education.add_run(edu_text).font.size = Pt(10)  # Adds education text

    if resume_data["projects"]:  # Checks if projects exist
        projects_title = doc.add_paragraph()  # Adds projects title paragraph
        projects_title.add_run("PROJECTS").bold = True  # Adds PROJECTS text
        projects_title.runs[0].font.size = Pt(12)  # Sets projects title font size
        projects = doc.add_paragraph()  # Adds projects paragraph
        projects.add_run("\n".join(resume_data["projects"])).font.size = Pt(10)  # Adds projects text

    if resume_data["certifications"]:  # Checks if certifications exist
        cert_title = doc.add_paragraph()  # Adds certifications title paragraph
        cert_title.add_run("CERTIFICATIONS").bold = True  # Adds CERTIFICATIONS text
        cert_title.runs[0].font.size = Pt(12)  # Sets certifications title font size
        certs = doc.add_paragraph()  # Adds certifications paragraph
        certs.add_run("\n".join(resume_data["certifications"])).font.size = Pt(10)  # Adds certifications text

    if resume_data["positions"]:  # Checks if positions exist
        pos_title = doc.add_paragraph()  # Adds positions title paragraph
        pos_title.add_run("EXTRACURRICULARS").bold = True  # Adds EXTRACURRICULARS text
        pos_title.runs[0].font.size = Pt(12)  # Sets positions title font size
        pos = doc.add_paragraph()  # Adds positions paragraph
        pos.add_run("\n".join(resume_data["positions"])).font.size = Pt(10)  # Adds positions text

    buffer = BytesIO()  # Creates in-memory buffer
    doc.save(buffer)  # Saves document to buffer
    buffer.seek(0)  # Resets buffer position
    if return_pdf:  # Checks if PDF output is requested
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:  # Creates temp Word file
            tmp.write(buffer.getvalue())  # Writes buffer to temp file
            tmp_path = tmp.name  # Gets temp file path
        pdf_path = tmp_path.replace(".docx", ".pdf")  # Creates PDF path
        docx2pdf.convert(tmp_path, pdf_path)  # Converts Word to PDF
        with open(pdf_path, "rb") as f:  # Opens PDF file
            pdf_buffer = BytesIO(f.read())  # Reads PDF into buffer
        os.remove(tmp_path)  # Deletes temp Word file
        os.remove(pdf_path)  # Deletes temp PDF file
        return pdf_buffer  # Returns PDF buffer
    return buffer  # Returns Word buffer

def generate_veteran_pdf_resume():  # Defines function for veteran PDF resume
    resume_data = st.session_state["resume_data"]  # Gets resume data
    buffer = BytesIO()  # Creates in-memory buffer
    c = canvas.Canvas(buffer, pagesize=letter)  # Creates PDF canvas
    width, height = letter  # Gets page dimensions

    c.setFillColor(HexColor("#0077b6"))  # Sets fill color for header
    c.rect(0, height - 50, width, 50, fill=1, stroke=0)  # Draws header rectangle

    c.setFont("Helvetica-Bold", 18)  # Sets font for name
    c.setFillColor(black)  # Sets text color to black
    c.drawString(50, height - 35, resume_data["personal_info"]["full_name"] or "Full Name Not Provided")  # Draws name
    c.setFont("Helvetica", 12)  # Sets font for contact info
    contact_parts = []  # Initializes contact parts list
    if resume_data["personal_info"]["address"]:  # Checks if address exists
        contact_parts.append(f"📍 {resume_data['personal_info']['address']}")  # Adds address
    else:
        contact_parts.append("📍 Address Not Provided")  # Adds placeholder if no address
    if resume_data["personal_info"]["phone"]:  # Checks if phone exists
        contact_parts.append(f"📞 {resume_data['personal_info']['phone']}")  # Adds phone
    else:
        contact_parts.append("📞 Phone Not Provided")  # Adds placeholder if no phone
    if resume_data["personal_info"]["email"]:  # Checks if email exists
        contact_parts.append(f"✉ {resume_data['personal_info']['email']}")  # Adds email
    else:
        contact_parts.append("✉ Email Not Provided")  # Adds placeholder if no email
    if resume_data["personal_info"]["linkedin"]:  # Checks if LinkedIn exists
        contact_parts.append(f"🔗 {resume_data['personal_info']['linkedin']}")  # Adds LinkedIn
    if resume_data["personal_info"]["github"]:  # Checks if GitHub exists
        contact_parts.append(f"🔗 {resume_data['personal_info']['github']}")  # Adds GitHub
    contact_info = "  |  ".join(contact_parts)  # Joins contact info
    c.drawString(50, height - 55, contact_info)  # Draws contact info

    y = height - 100  # Sets initial y position

    def add_section(title, content, y_pos):  # Defines function to add PDF section
        c.setFont("Helvetica-Bold", 14)  # Sets font for section title
        c.setFillColor(HexColor("#0077b6"))  # Sets title color
        c.drawString(50, y_pos, title)  # Draws section title
        c.setFillColor(black)  # Sets text color to black
        c.setFont("Helvetica", 10)  # Sets font for content
        y_pos -= 20  # Adjusts y position
        if isinstance(content, list):  # Checks if content is a list
            for line in content:  # Loops through content lines
                if line:  # Checks if line is non-empty
                    c.drawString(50, y_pos, line)  # Draws content line
                    y_pos -= 15  # Adjusts y position
        else:
            c.drawString(50, y_pos, content)  # Draws single content line
            y_pos -= 15  # Adjusts y position
        return y_pos - 20  # Returns new y position

    y = add_section("Summary", resume_data["summary"] or "Summary Not Provided", y)  # Adds Summary section
    education_content = resume_data["qualifications"] if resume_data["qualifications"] else ["Education Not Provided"]  # Gets education content
    y = add_section("Education", education_content, y)  # Adds Education section
    employment_content = [f"{exp.get('job_title', 'Title Not Provided')} - {exp.get('company', 'Company Not Provided')} ({exp.get('dates', 'Dates Not Provided')})" for exp in resume_data["experience"]] if resume_data["experience"] else ["Employment Not Provided"]  # Gets employment content
    y = add_section("Employment", employment_content, y)  # Adds Employment section
    if resume_data["certifications"]:  # Checks if certifications exist
        y = add_section("Certifications", resume_data["certifications"], y)  # Adds Certifications section
    achievements_content = [f"- {exp.get('achievements', '')}" for exp in resume_data["experience"] if exp.get('achievements')] if any(exp.get('achievements') for exp in resume_data["experience"]) else ["Achievements Not Provided"]  # Gets achievements content
    y = add_section("Achievements", achievements_content, y)  # Adds Achievements section
    if resume_data["positions"]:  # Checks if positions exist
        y = add_section("Extracurriculars", resume_data["positions"], y)  # Adds Extracurriculars section
    skills_content = ", ".join(resume_data["skills"]) if resume_data["skills"] else "Skills Not Provided"  # Gets skills content
    y = add_section("Skills", skills_content, y)  # Adds Skills section
    languages_content = "English"  # Sets languages content
    y = add_section("Languages", languages_content, y)  # Adds Languages section

    c.save()  # Saves PDF
    buffer.seek(0)  # Resets buffer position
    return buffer  # Returns PDF buffer

def generate_interview_questions():  # Defines function to generate interview questions
    resume_data = st.session_state["resume_data"]  # Gets resume data
    resume_content = (  # Formats resume content for prompt
        f"Personal Info: {resume_data['personal_info']}\n"
        f"Summary: {resume_data['summary']}\n"
        f"Experience: {resume_data['experience']}\n"
        f"Projects: {resume_data['projects']}\n"
        f"Qualifications: {resume_data['qualifications']}\n"
        f"Skills: {resume_data['skills']}\n"
        f"Certifications: {resume_data['certifications']}\n"
        f"Positions of Responsibility: {resume_data['positions']}"
    )

    prompt = f"""  # Defines prompt for interview questions
    Based on the following resume data, generate 20 specific interview questions tailored to the candidate's experience, skills, projects, education, certifications, and achievements. Ensure the questions are relevant, insightful, and encourage detailed responses about their background. Number each question from 1 to 20 and list them clearly.

    Resume Data:
    {resume_content}
    """

    response = get_gemini_response(prompt)  # Gets response from Gemini
    if response:  # Checks if response exists
        return response  # Returns interview questions
    else:
        return "Failed to generate interview questions."  # Returns error message

def preview_pdf_scrollable(pdf_buffer):  # Defines function to preview PDF
    pdf_buffer.seek(0)  # Resets buffer position
    base64_pdf = base64.b64encode(pdf_buffer.read()).decode('utf-8')  # Encodes PDF to base64
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600px" type="application/pdf" title="Resume Preview"></iframe>'  # Creates iframe for PDF
    st.markdown(pdf_display, unsafe_allow_html=True)  # Displays PDF in Streamlit

def create_download_link(file_path_or_buffer, file_name, link_text):  # Defines function to create download link
    if isinstance(file_path_or_buffer, str):  # Checks if input is a file path
        with open(file_path_or_buffer, "rb") as f:  # Opens file
            file_content = f.read()  # Reads file content
    else:
        file_content = file_path_or_buffer.getvalue()  # Gets buffer content
    b64 = base64.b64encode(file_content).decode()  # Encodes content to base64
    mime_type = "application/pdf" if file_name.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # Sets MIME type
    return f'<a href="data:{mime_type};base64,{b64}" download="{file_name}" target="_blank">{link_text}</a>'  # Returns HTML download link

st.markdown("""  # Defines custom CSS for styling
    <style>
    .stTextInput, .stTextArea {
        margin-bottom: 20px;
        padding: 10px;
        width: 100%;
    }
    .stRadio > div {
        margin-bottom: 20px;
    }
    .stButton > button {
        margin-top: 10px;
        margin-bottom: 20px;
        padding: 10px 20px;
    }
    .stExpander {
        margin-bottom: 20px;
    }
    .question-section {
        font-size: 18px;
        line-height: 1.5;
        margin-bottom: 20px;
    }
    .language-instruction {
        font-size: 14px;
        color: #888;
        margin-bottom: 20px;
    }
    .personal-info-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 20px;
    }
    button:focus {
        outline: 2px solid #0078d4;
    }
    input:focus, textarea:focus {
        outline: 2px solid #0078d4;
    }
    </style>
""", unsafe_allow_html=True)

with st.container():  # Creates a container for layout
    col1, col2 = st.columns([1, 4])  # Splits layout into two columns
    with col1:  # First column
        st.image("/Users/vaidhav/Downloads/Ruby Williams.png", width=100, caption="ResWhisper Icon")  # Displays app icon
    with col2:  # Second column
        st.title("ResuWhisper AI")  # Displays app title
        st.write("Create a professional resume with your voice")  # Displays subtitle

if st.session_state["page"] != "login" and st.session_state["page"] != "signup" and st.session_state["page"] != "welcome":  # Checks if not on login/signup/welcome pages
    progress_value = 0  # Initializes progress value
    total_questions = len(questions)  # Gets total number of questions
    if st.session_state["page"] == "language_selection":  # Checks if on language selection page
        progress_value = 10  # Sets progress to 10%
    elif st.session_state["page"] == "consent":  # Checks if on consent page
        progress_value = 15  # Sets progress to 15%
    elif st.session_state["page"] == "resume_template":  # Checks if on template page
        progress_value = 20  # Sets progress to 20%
    elif st.session_state["page"] == "questions":  # Checks if on questions page
        progress_value = 20 + (st.session_state["current_question_index"] / total_questions * 60)  # Calculates progress
    elif st.session_state["page"] == "preview":  # Checks if on preview page
        progress_value = 100  # Sets progress to 100%
    st.progress(progress_value / 100, text=f"Progress: {int(progress_value)}%")  # Displays progress bar

if st.session_state["page"] == "login":  # Checks if on login page
    st.title("Login")  # Displays login title
    username = st.text_input("Username")  # Creates username input field
    password = st.text_input("Password", type="password")  # Creates password input field
    if st.button("Login"):  # Creates login button
        if username in st.session_state["users_db"] and st.session_state["users_db"][username]["password"] == password:  # Checks credentials
            st.session_state["authenticated"] = True  # Sets authenticated flag
            st.session_state["username"] = username  # Sets username
            st.session_state["resume_data"] = st.session_state["users_db"][username].get("data", st.session_state["resume_data"])  # Loads user data
            st.session_state["page"] = "welcome"  # Moves to welcome page
            st.success("Logged in successfully!")  # Displays success message
            st.rerun()  # Reruns app
        else:
            st.error("Invalid username or password")  # Displays error message
    if st.button("Sign Up Instead"):  # Creates signup button
        st.session_state["page"] = "signup"  # Moves to signup page
        st.rerun()  # Reruns app

elif st.session_state["page"] == "signup":  # Checks if on signup page
    st.title("Sign Up")  # Displays signup title
    new_username = st.text_input("New Username")  # Creates username input field
    new_password = st.text_input("New Password", type="password")  # Creates password input field
    confirm_password = st.text_input("Confirm Password", type="password")  # Creates confirm password field
    if st.button("Sign Up"):  # Creates signup button
        if new_username in st.session_state["users_db"]:  # Checks if username exists
            st.error("Username already exists")  # Displays error
        elif new_password != confirm_password:  # Checks if passwords match
            st.error("Passwords do not match")  # Displays error
        elif new_username and new_password:  # Checks if fields are filled
            st.session_state["users_db"][new_username] = {"password": new_password, "data": st.session_state["resume_data"]}  # Adds new user
            st.session_state["authenticated"] = True  # Sets authenticated flag
            st.session_state["username"] = new_username  # Sets username
            st.session_state["page"] = "welcome"  # Moves to welcome page
            st.success("Signed up successfully!")  # Displays success message
            st.rerun()  # Reruns app
        else:
            st.error("Please fill all fields")  # Displays error
    if st.button("Back to Login"):  # Creates back button
        st.session_state["page"] = "login"  # Moves to login page
        st.rerun()  # Reruns app

elif st.session_state["page"] == "welcome" and st.session_state["authenticated"]:  # Checks if on welcome page and authenticated
    st.markdown("""  # Displays welcome message
    ### Build Your Professional Resume with Just Your Voice!
    
    ResuWhisper AI uses advanced AI to help you create a professional resume simply by answering a few questions. 
    Speak naturally, and we'll convert your responses into a polished, well-formatted resume.
    
    **How it works:**
    1. Select a language for your responses
    2. Choose a resume template that fits your career stage
    3. Answer a series of questions using your voice or by uploading audio
    4. The AI processes your responses into professional resume content
    5. Preview and download your completed resume
    
    **Key Features:**
    - Voice-to-text conversion for easy input
    - AI-powered content formatting and enhancement
    - Multiple professional resume templates
    - Download in PDF or Word format
    """)
    
    if st.button("🚀 Get Started", key="welcome_start"):  # Creates get started button
        st.session_state["page"] = "language_selection"  # Moves to language selection
        st.rerun()  # Reruns app

elif st.session_state["page"] == "language_selection":  # Checks if on language selection page
    st.title("🌍 Choose Your Preferred Language")  # Displays title
    st.write("Select the language you’ll use to answer the questions (will be transcribed to English):")  # Displays instruction
    
    cols = st.columns(4)  # Creates 4 columns for language buttons
    for i, lang in enumerate(languages):  # Loops through languages
        with cols[i % 4]:  # Uses modulo to distribute languages
            if st.button(lang, key=f"lang_{lang}", help=f"Select {lang} as your language"):  # Creates language button
                st.session_state["selected_language"] = lang  # Sets selected language
                if lang != "English":  # Checks if not English
                    translate_questions(lang)  # Translates questions
                st.rerun()  # Reruns app
    
    col1, col2 = st.columns([1, 1])  # Creates two columns for navigation
    with col1:  # First column
        if st.button("← Back to Welcome", key="lang_back"):  # Creates back button
            st.session_state["page"] = "welcome"  # Moves to welcome page
            st.rerun()  # Reruns app
    with col2:  # Second column
        if st.session_state["selected_language"]:  # Checks if language is selected
            if st.button("Next →", key="lang_next"):  # Creates next button
                st.session_state["page"] = "consent"  # Moves to consent page
                st.rerun()  # Reruns app
        else:
            st.info("Please select a language")  # Displays info message

elif st.session_state["page"] == "consent":  # Checks if on consent page
    st.title("📜 User Consent")  # Displays title
    st.markdown("""  # Displays consent message
    Before we proceed, please give us your consent to use your voice input for processing your resume.
    
    By checking the box below, you agree to allow this application to record and process your voice responses.
    """)
    st.session_state["consent_given"] = st.checkbox("I agree to the terms and conditions.", key="consent_checkbox")  # Creates consent checkbox
    col1, col2 = st.columns([1, 1])  # Creates two columns for navigation
    with col1:  # First column
        if st.button("← Back to Language Selection", key="consent_back"):  # Creates back button
            st.session_state["page"] = "language_selection"  # Moves to language selection
            st.rerun()  # Reruns app
    with col2:  # Second column
        if st.session_state["consent_given"]:  # Checks if consent is given
            if st.button("Proceed →", key="consent_next"):  # Creates proceed button
                st.session_state["page"] = "resume_template"  # Moves to template page
                st.rerun()  # Reruns app
        else:
            st.info("Please give your consent to proceed.")  # Displays info message

elif st.session_state["page"] == "resume_template":  # Checks if on template page
    st.title("📄 Choose Your Resume Template")  # Displays title
    st.write("Select a template that best matches your career stage:")  # Displays instruction
    
    col1, col2 = st.columns(2)  # Creates two columns for templates
    
    with col1:  # First column
        st.subheader("Entry Level / Student Template")  # Displays subheader
        st.image("https://jofibostorage.blob.core.windows.net/blog/Fresher-reverse.chronological.png", width=200, caption="Entry Level Template")  # Displays fresher template image
        st.write(resume_templates["Fresher"])  # Displays fresher description
        if st.button("Select Entry Level Template", key="fresher_template"):  # Creates fresher button
            st.session_state["resume_template"] = "Fresher"  # Sets template to Fresher
            st.session_state["page"] = "questions"  # Moves to questions page
            st.rerun()  # Reruns app
        
        st.subheader("Executive / Senior Level Template")  # Displays subheader
        st.image("https://cdn-blog.novoresume.com/articles/executive-resume-example/executive-resume-sample.png", width=200, caption="Executive Template")  # Displays veteran template image
        st.write(resume_templates["Veteran"])  # Displays veteran description
        if st.button("Select Executive Template", key="veteran_template"):  # Creates veteran button
            st.session_state["resume_template"] = "Veteran"  # Sets template to Veteran
            st.session_state["page"] = "questions"  # Moves to questions page
            st.rerun()  # Reruns app
    
    with col2:  # Second column
        st.subheader("Mid-Career Professional Template")  # Displays subheader
        st.image("https://www.resumebuilder.com/wp-content/uploads/2020/12/Sales-Executive-Resume-Example-Banner-Image.png", width=200, caption="Mid-Career Template")  # Displays intermediate template image
        st.write(resume_templates["Intermediate"])  # Displays intermediate description
        if st.button("Select Mid-Career Template", key="intermediate_template"):  # Creates intermediate button
            st.session_state["resume_template"] = "Intermediate"  # Sets template to Intermediate
            st.session_state["page"] = "questions"  # Moves to questions page
            st.rerun()  # Reruns app
    
    if st.button("← Back to Consent", key="template_back"):  # Creates back button
        st.session_state["page"] = "consent"  # Moves to consent page
        st.rerun()  # Reruns app

elif st.session_state["page"] == "questions":  # Checks if on questions page
    st.title("Resume Information Questionnaire")  # Displays title
    
    current_q = st.session_state["current_question_index"] + 1  # Calculates current question number
    total_q = len(questions)  # Gets total questions
    st.write(f"Question {current_q} of {total_q}")  # Displays question progress
    
    st.header(section_headers[st.session_state["current_question_index"]])  # Displays section header
    
    col1, col2 = st.columns([2, 3])  # Creates two columns for layout

    with col1:  # First column
        current_index = st.session_state["current_question_index"]  # Gets current question index
        current_question = questions[current_index]  # Gets current question
        st.markdown(f"### {current_question}")  # Displays current question
        if st.session_state["selected_language"] != "English" and st.session_state["selected_language"] in st.session_state["translated_questions"]:  # Checks if translated
            translated_question = st.session_state["translated_questions"][st.session_state["selected_language"]][current_index]  # Gets translated question
            st.markdown(f"### {translated_question}")  # Displays translated question
        st.markdown(f'<div class="language-instruction">Please answer in {st.session_state["selected_language"]} (will be transcribed to English).</div>', unsafe_allow_html=True)  # Displays language instruction

        input_option = st.radio("Choose Input Method:", ["Record Audio", "Upload Audio", "Text"], horizontal=True, key=f"input_method_{current_index}")  # Creates input method radio buttons

        if input_option == "Record Audio":  # Checks if recording selected
            if st.session_state["audio_file"] is None:  # Checks if audio file is not set
                st.session_state["audio_file"] = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name  # Creates temp audio file

            col_a, col_b = st.columns(2)  # Creates two columns for recording buttons
            with col_a:  # First column
                if not st.session_state.get("recording_state", False):  # Checks if not recording
                    if st.button("🎤 Start Recording", key=f"start_rec_{current_index}"):  # Creates start button
                        st.session_state["recording_state"] = True  # Sets recording state
                        stop_event = threading.Event()  # Creates stop event
                        st.session_state["stop_event"] = stop_event  # Stores stop event
                        recording_thread = threading.Thread(target=record_audio, args=(st.session_state["audio_file"], stop_event), daemon=True)  # Creates recording thread
                        st.session_state["recording_thread"] = recording_thread  # Stores thread
                        recording_thread.start()  # Starts recording
                        st.rerun()  # Reruns app
            with col_b:  # Second column
                if st.session_state.get("recording_state", False):  # Checks if recording
                    st.warning("🔴 Recording in progress... speak clearly!")  # Displays warning
                                        if st.button("⏹ Stop Recording", key=f"stop_rec_{current_index}"):  # Creates stop recording button
                        st.session_state["recording_state"] = False  # Sets recording state to False
                        if "stop_event" in st.session_state:  # Checks if stop event exists
                            st.session_state["stop_event"].set()  # Signals to stop recording
                        if "recording_thread" in st.session_state and st.session_state["recording_thread"].is_alive():  # Checks if thread is running
                            st.session_state["recording_thread"].join(timeout=5)  # Waits for thread to finish
                        if os.path.exists(st.session_state["audio_file"]) and os.path.getsize(st.session_state["audio_file"]) > 0:  # Checks if audio file exists and has data
                            st.success(f"✅ Recording saved to {st.session_state['audio_file']}!")  # Displays success message
                            with st.spinner("Transcribing..."):  # Shows transcription spinner
                                progress = st.progress(0)  # Initializes progress bar
                                prompt = f"Transcribe the audio to English, assuming it is spoken in {st.session_state['selected_language']}."  # Creates transcription prompt
                                response = get_gemini_response(prompt, st.session_state["audio_file"])  # Gets transcription from Gemini
                                progress.progress(100)  # Completes progress bar
                                if response:  # Checks if transcription succeeded
                                    st.session_state["current_response"] = response  # Stores transcribed response
                                    update_resume_data(current_index, response)  # Updates resume data with response
                                else:
                                    st.error("Transcription failed after recording")  # Displays error if transcription fails
                        else:
                            st.error("Audio file was not created or is empty")  # Displays error if file is missing/empty
                        if "recording_thread" in st.session_state:  # Checks if thread exists
                            del st.session_state["recording_thread"]  # Deletes thread from session state
                        if "stop_event" in st.session_state:  # Checks if stop event exists
                            del st.session_state["stop_event"]  # Deletes stop event from session state
                        st.rerun()  # Reruns app to refresh UI

        elif input_option == "Upload Audio":  # Checks if upload audio is selected
            uploaded_file = st.file_uploader("Upload an audio file (MP3, WAV)", type=["mp3", "wav"], key=f"upload_{current_index}")  # Creates file uploader
            if uploaded_file:  # Checks if file is uploaded
                if st.session_state["audio_file"] is None:  # Checks if audio file path is not set
                    st.session_state["audio_file"] = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name  # Creates temp file path
                with open(st.session_state["audio_file"], "wb") as f:  # Opens temp file for writing
                    f.write(uploaded_file.getvalue())  # Writes uploaded file to temp file
                if os.path.exists(st.session_state["audio_file"]) and os.path.getsize(st.session_state["audio_file"]) > 0:  # Checks if file exists and has data
                    st.success("✅ Audio file uploaded!")  # Displays success message
                    if current_index not in st.session_state["transcribed_once"]:  # Checks if not already transcribed
                        with st.spinner("Transcribing..."):  # Shows transcription spinner
                            progress = st.progress(0)  # Initializes progress bar
                            prompt = f"Transcribe the audio to English, assuming it is spoken in {st.session_state['selected_language']}."  # Creates transcription prompt
                            response = get_gemini_response(prompt, st.session_state["audio_file"])  # Gets transcription
                            progress.progress(100)  # Completes progress bar
                            if response:  # Checks if transcription succeeded
                                st.session_state["current_response"] = response  # Stores transcribed response
                                update_resume_data(current_index, response)  # Updates resume data
                                st.session_state["transcribed_once"][current_index] = True  # Marks as transcribed
                            else:
                                st.error("Transcription failed")  # Displays error if transcription fails
                    else:
                        st.info("Audio already transcribed. Edit your response below if needed.")  # Informs user if already transcribed
                else:
                    st.error("Failed to save uploaded file")  # Displays error if file save fails

        elif input_option == "Text":  # Checks if text input is selected
            text_input = st.text_area("Type your response here:", height=200, key=f"text_{current_index}")  # Creates text area for input
            if st.button("Submit Text", key=f"submit_text_{current_index}"):  # Creates submit button
                if text_input:  # Checks if text is entered
                    st.session_state["current_response"] = text_input  # Stores text response
                    update_resume_data(current_index, text_input)  # Updates resume data
                else:
                    st.error("Please enter some text.")  # Displays error if text is empty

        if st.session_state["current_response"]:  # Checks if there’s a current response
            st.subheader("Your Response:")  # Displays response header
            edited_response = st.text_area("Edit your response if needed:", st.session_state["current_response"], height=200, key=f"trans_{current_index}")  # Creates editable text area
            can_proceed = True  # Initializes proceed flag
            if current_index not in [3, 6, 7]:  # Checks if question is mandatory (not projects, certifications, extracurriculars)
                if not edited_response.strip():  # Checks if response is empty
                    st.error("This question is mandatory. Please provide a response.")  # Displays error
                    can_proceed = False  # Disables proceeding
            if st.button("✅ Process Response and Continue", key=f"process_{current_index}", disabled=not can_proceed):  # Creates process button
                st.session_state["current_response"] = edited_response  # Updates current response
                st.session_state["responses"][current_index] = edited_response  # Stores response in responses dict
                if current_index < len(questions) - 1:  # Checks if not last question
                    st.session_state["current_question_index"] += 1  # Moves to next question
                    if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):  # Checks if audio file exists
                        os.remove(st.session_state["audio_file"])  # Deletes audio file
                    st.session_state["audio_file"] = None  # Resets audio file
                    st.session_state["current_response"] = None  # Resets current response
                    st.rerun()  # Reruns app
                else:
                    st.session_state["page"] = "preview"  # Moves to preview page if last question
                    st.rerun()  # Reruns app

        col_nav1, col_nav2, col_nav3 = st.columns([1, 1, 1])  # Creates three columns for navigation
        with col_nav1:  # First column
            if st.session_state["current_question_index"] > 0:  # Checks if not first question
                if st.button("← Previous Question", key=f"prev_{current_index}"):  # Creates previous button
                    st.session_state["current_question_index"] -= 1  # Moves to previous question
                    st.session_state["current_response"] = None  # Resets current response
                    if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):  # Checks if audio file exists
                        os.remove(st.session_state["audio_file"])  # Deletes audio file
                    st.session_state["audio_file"] = None  # Resets audio file
                    st.rerun()  # Reruns app
        with col_nav3:  # Third column
            if current_index in [3, 6, 7]:  # Checks if question is optional (projects, certifications, extracurriculars)
                if st.button("Skip This Question →", key=f"skip_{current_index}"):  # Creates skip button
                    if current_index < len(questions) - 1:  # Checks if not last question
                        st.session_state["current_question_index"] += 1  # Moves to next question
                        st.session_state["current_response"] = None  # Resets current response
                        if st.session_state["audio_file"] and os.path.exists(st.session_state["audio_file"]):  # Checks if audio file exists
                            os.remove(st.session_state["audio_file"])  # Deletes audio file
                        st.session_state["audio_file"] = None  # Resets audio file
                        st.rerun()  # Reruns app
                    else:
                        st.session_state["page"] = "preview"  # Moves to preview page if last question
                        st.rerun()  # Reruns app

    with col2:  # Second column
        st.subheader("Live Resume Editor")  # Displays editor header
        st.markdown("**You can edit the Live Resume Editor if you are not satisfied. Type it or re-record yourself, it all works!**")  # Displays instruction
        resume_data = st.session_state["resume_data"]  # Gets resume data

        with st.expander("📋 Personal Information", expanded=True):  # Creates personal info expander
            st.markdown('<div class="personal-info-grid">', unsafe_allow_html=True)  # Starts CSS grid for layout
            col_a, col_b = st.columns(2)  # Creates two columns
            with col_a:  # First column
                new_full_name = st.text_input("Full Name", resume_data["personal_info"]["full_name"], key="personal_full_name")  # Creates full name input
                resume_data["personal_info"]["full_name"] = new_full_name  # Updates full name
            with col_b:  # Second column
                new_degree = st.text_input("Degree", resume_data["personal_info"]["degree"], key="personal_degree")  # Creates degree input
                resume_data["personal_info"]["degree"] = new_degree  # Updates degree
            col_c, col_d = st.columns(2)  # Creates two more columns
            with col_c:  # First column
                new_phone = st.text_input("Phone", resume_data["personal_info"]["phone"], key="personal_phone")  # Creates phone input
                resume_data["personal_info"]["phone"] = new_phone  # Updates phone
            with col_d:  # Second column
                new_email = st.text_input("Email", resume_data["personal_info"]["email"], key="personal_email")  # Creates email input
                resume_data["personal_info"]["email"] = new_email  # Updates email
            col_e, col_f = st.columns(2)  # Creates two more columns
            with col_e:  # First column
                new_linkedin = st.text_input("LinkedIn (Optional)", resume_data["personal_info"]["linkedin"], key="personal_linkedin")  # Creates LinkedIn input
                resume_data["personal_info"]["linkedin"] = new_linkedin  # Updates LinkedIn
            with col_f:  # Second column
                new_github = st.text_input("GitHub (Optional)", resume_data["personal_info"]["github"], key="personal_github")  # Creates GitHub input
                resume_data["personal_info"]["github"] = new_github  # Updates GitHub
            col_g, _ = st.columns(2)  # Creates two more columns (only uses first)
            with col_g:  # First column
                new_address = st.text_input("Address", resume_data["personal_info"]["address"], key="personal_address")  # Creates address input
                resume_data["personal_info"]["address"] = new_address  # Updates address
            st.markdown('</div>', unsafe_allow_html=True)  # Ends CSS grid

            if not all([resume_data["personal_info"]["full_name"], resume_data["personal_info"]["degree"],  # Checks if mandatory fields are filled
                        resume_data["personal_info"]["phone"], resume_data["personal_info"]["email"],
                        resume_data["personal_info"]["address"]]):
                st.error("All fields except LinkedIn and GitHub are mandatory.")  # Displays error if fields are missing

        with st.expander("💼 Summary", expanded=True):  # Creates summary expander
            new_summary = st.text_area("Summary", resume_data["summary"], height=100, key=f"summary_{current_index}")  # Creates summary text area
            resume_data["summary"] = new_summary  # Updates summary
            if current_index >= 1 and not new_summary.strip():  # Checks if summary is mandatory and empty
                st.error("Summary is mandatory.")  # Displays error

        if st.session_state["current_question_index"] >= 2:  # Checks if experience section is accessible
            with st.expander("👔 Experience", expanded=True):  # Creates experience expander
                if not resume_data["experience"]:  # Checks if experience is empty
                    resume_data["experience"] = [{"job_title": "", "company": "", "dates": "", "responsibilities": "", "achievements": ""}]  # Initializes empty experience
                for i, exp in enumerate(resume_data["experience"]):  # Loops through experience entries
                    st.write(f"Experience {i+1}")  # Displays experience number
                    exp["job_title"] = st.text_input(f"Job Title {i+1}", exp.get("job_title", ""), key=f"exp_title_{i}")  # Creates job title input
                    exp["company"] = st.text_input(f"Company {i+1}", exp.get("company", ""), key=f"exp_company_{i}")  # Creates company input
                    exp["dates"] = st.text_input(f"Dates {i+1}", exp.get("dates", ""), key=f"exp_dates_{i}")  # Creates dates input
                    exp["responsibilities"] = st.text_area(f"Responsibilities {i+1}", exp.get("responsibilities", ""), key=f"exp_resp_{i}")  # Creates responsibilities input
                    exp["achievements"] = st.text_area(f"Achievements {i+1}", exp.get("achievements", ""), key=f"exp_achieve_{i}")  # Creates achievements input
                if not any(exp["job_title"].strip() for exp in resume_data["experience"]):  # Checks if experience is empty
                    st.error("Experience is mandatory.")  # Displays error

        if st.session_state["current_question_index"] >= 3:  # Checks if projects section is accessible
            with st.expander("🚀 Projects (Optional)", expanded=True):  # Creates projects expander
                if not resume_data["projects"]:  # Checks if projects are empty
                    resume_data["projects"] = [""]  # Initializes empty project
                for i, proj in enumerate(resume_data["projects"]):  # Loops through projects
                    new_proj = st.text_input(f"Project {i+1}", proj, key=f"proj_{i}")  # Creates project input
                    resume_data["projects"][i] = new_proj  # Updates project

        if st.session_state["current_question_index"] >= 4:  # Checks if qualifications section is accessible
            with st.expander("🎓 Qualifications", expanded=True):  # Creates qualifications expander
                if not resume_data["qualifications"]:  # Checks if qualifications are empty
                    resume_data["qualifications"] = [""]  # Initializes empty qualification
                for i, qual in enumerate(resume_data["qualifications"]):  # Loops through qualifications
                    new_qual = st.text_input(f"Qualification {i+1}", qual, key=f"qual_{i}")  # Creates qualification input
                    resume_data["qualifications"][i] = new_qual  # Updates qualification
                if not any(qual.strip() for qual in resume_data["qualifications"]):  # Checks if qualifications are empty
                    st.error("Qualifications are mandatory.")  # Displays error

        if st.session_state["current_question_index"] >= 5:  # Checks if skills section is accessible
            with st.expander("🛠️ Skills", expanded=True):  # Creates skills expander
                if not resume_data["skills"]:  # Checks if skills are empty
                    resume_data["skills"] = [""]  # Initializes empty skill
                for i, skill in enumerate(resume_data["skills"]):  # Loops through skills
                    new_skill = st.text_input(f"Skill {i+1}", skill, key=f"skill_{i}")  # Creates skill input
                    resume_data["skills"][i] = new_skill  # Updates skill
                if not any(skill.strip() for skill in resume_data["skills"]):  # Checks if skills are empty
                    st.error("Skills are mandatory.")  # Displays error

        if st.session_state["current_question_index"] >= 6:  # Checks if certifications section is accessible
            with st.expander("🏆 Certifications (Optional)", expanded=True):  # Creates certifications expander
                if not resume_data["certifications"]:  # Checks if certifications are empty
                    resume_data["certifications"] = [""]  # Initializes empty certification
                for i, cert in enumerate(resume_data["certifications"]):  # Loops through certifications
                    new_cert = st.text_input(f"Certification {i+1}", cert, key=f"cert_{i}")  # Creates certification input
                    resume_data["certifications"][i] = new_cert  # Updates certification

        if st.session_state["current_question_index"] >= 7:  # Checks if positions section is accessible
            with st.expander("🌟 Positions of Responsibility (Optional)", expanded=True):  # Creates positions expander
                if not resume_data["positions"]:  # Checks if positions are empty
                    resume_data["positions"] = [""]  # Initializes empty position
                for i, pos in enumerate(resume_data["positions"]):  # Loops through positions
                    new_pos = st.text_input(f"Position {i+1}", pos, key=f"pos_{i}")  # Creates position input
                    resume_data["positions"][i] = new_pos  # Updates position

        st.session_state["resume_data"] = resume_data  # Updates resume data in session state
        if st.session_state["username"]:  # Checks if user is logged in
            st.session_state["users_db"][st.session_state["username"]]["data"] = resume_data  # Updates user data in database

elif st.session_state["page"] == "preview":  # Checks if on preview page
    st.title("Your Resume Preview")  # Displays title
    
    st.subheader("Interview Questions Based on Your Resume")  # Displays subheader
    with st.spinner("Generating interview questions..."):  # Shows spinner while generating
        interview_questions = generate_interview_questions()  # Generates interview questions
        st.markdown(interview_questions)  # Displays questions
        st.download_button("Download Questions", interview_questions, "interview_questions.txt", key="download_questions")  # Creates download button for questions

    st.subheader("PDF Preview")  # Displays subheader
    with st.spinner("Generating PDF Preview..."):  # Shows spinner while generating
        if st.session_state["resume_template"] == "Fresher":  # Checks if Fresher template
            pdf_buffer = generate_word_resume(return_pdf=True)  # Generates PDF for Fresher
            preview_pdf_scrollable(pdf_buffer)  # Previews PDF
        elif st.session_state["resume_template"] == "Intermediate":  # Checks if Intermediate template
            pdf_buffer = generate_intermediate_word_resume(return_pdf=True)  # Generates PDF for Intermediate
            preview_pdf_scrollable(pdf_buffer)  # Previews PDF
        elif st.session_state["resume_template"] == "Veteran":  # Checks if Veteran template
            pdf_buffer = generate_veteran_pdf_resume()  # Generates PDF for Veteran
            preview_pdf_scrollable(pdf_buffer)  # Previews PDF
        else:
            st.error("Please select a resume template to preview.")  # Displays error if no template selected

    st.subheader("Download Your Resume")  # Displays subheader
    file_format = st.selectbox("Choose format:", ["PDF", "Word"], key="download_format")  # Creates format selector
    
    if st.button("📥 Generate Resume for Download", key="download_button"):  # Creates download button
        if st.session_state["resume_template"] == "Fresher":  # Checks if Fresher template
            if file_format == "PDF":  # Checks if PDF format
                with st.spinner("Generating PDF..."):  # Shows spinner
                    pdf_buffer = generate_word_resume(return_pdf=True)  # Generates PDF
                    download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")  # Creates download link
                    st.markdown(download_link, unsafe_allow_html=True)  # Displays download link
            elif file_format == "Word":  # Checks if Word format
                with st.spinner("Generating Word document..."):  # Shows spinner
                    word_buffer = generate_word_resume(return_pdf=False)  # Generates Word doc
                    download_link = create_download_link(word_buffer, "resume.docx", "Download Word Document")  # Creates download link
                    st.markdown(download_link, unsafe_allow_html=True)  # Displays download link
        elif st.session_state["resume_template"] == "Intermediate":  # Checks if Intermediate template
            if file_format == "PDF":  # Checks if PDF format
                with st.spinner("Generating PDF..."):  # Shows spinner
                    pdf_buffer = generate_intermediate_word_resume(return_pdf=True)  # Generates PDF
                    download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")  # Creates download link
                    st.markdown(download_link, unsafe_allow_html=True)  # Displays download link
            elif file_format == "Word":  # Checks if Word format
                with st.spinner("Generating Word document..."):  # Shows spinner
                    word_buffer = generate_intermediate_word_resume(return_pdf=False)  # Generates Word doc
                    download_link = create_download_link(word_buffer, "resume.docx", "Download Word Document")  # Creates download link
                    st.markdown(download_link, unsafe_allow_html=True)  # Displays download link
        elif st.session_state["resume_template"] == "Veteran":  # Checks if Veteran template
            with st.spinner("Generating PDF..."):  # Shows spinner
                pdf_buffer = generate_veteran_pdf_resume()  # Generates PDF
                download_link = create_download_link(pdf_buffer, "resume.pdf", "Download PDF")  # Creates download link
                st.markdown(download_link, unsafe_allow_html=True)  # Displays download link
        else:
            st.error("Please select a resume template to download.")  # Displays error if no template selected

    st.markdown("---")  # Adds horizontal line
    if st.button("← Back to Questionnaire", key="preview_back"):  # Creates back button
        st.session_state["page"] = "questions"  # Moves to questions page
        st.rerun()  # Reruns app
    
    if st.button("🔄 Start Over", key="start_over"):  # Creates start over button
        for key in list(st.session_state.keys()):  # Loops through session state keys
            if key != "consent_given" and key != "authenticated" and key != "username" and key != "users_db":  # Excludes certain keys
                if key == "page":  # Checks if page key
                    st.session_state[key] = "welcome"  # Resets to welcome page
                elif key == "responses":  # Checks if responses key
                    st.session_state[key] = {}  # Resets responses
                elif key == "current_question_index":  # Checks if question index key
                    st.session_state[key] = 0  # Resets question index
                elif key == "resume_data":  # Checks if resume data key
                    st.session_state[key] = {  # Resets resume data
                        "personal_info": {
                            "full_name": "",
                            "degree": "",
                            "phone": "",
                            "email": "",
                            "linkedin": "",
                            "github": "",
                            "address": ""
                        },
                        "summary": "",
                        "qualifications": [],
                        "certifications": [],
                        "skills": [],
                        "experience": [],
                        "projects": [],
                        "positions": []
                    }
                elif key in ["resume_template", "current_response", "selected_language"]:  # Checks if specific keys
                    st.session_state[key] = None  # Resets to None
        st.rerun()  # Reruns app

st.markdown("---")  # Adds horizontal line
st.markdown("""  # Displays footer
<div style="text-align: center;">
    <p>© 2025 Resume Builder AI | Powered by CSSTUV</p>
</div>
""", unsafe_allow_html=True)